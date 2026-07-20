"""
Entry Requirements Checker — lightweight web tool
=================================================
Validates a course page's Entry Requirements against its Qualification
Specification document (and the tracker Excel), using AI comparison.

Workflow
--------
1. Upload the tracker Excel (Course Name / Course URL / Specification
   Document URL / Level / Type — Level & Type are parsed from the course
   name when the columns are absent).
2. Every unique Qualification Specification document is processed ONCE:
   its Entry Requirements are extracted and cached in SQLite. Courses that
   share the same spec URL reuse the cached extraction.
3. Pick a course → Run Check: the course page's Entry Requirements are
   extracted live, compared (AI) against the cached spec + the Excel value.
4. A Pass/Fail report is shown and downloadable as PDF.
5. The original Content Quality Review remains available for proofreading.
6. A separate Course Overview Quality page can load a selected tracker course's
   webpage overview, accept pasted text, or extract the Course Overview from
   TXT/DOCX/PDF content for an editorial quality evaluation and improved copy.

Persistence
-----------
Everything the user uploads (courses, cached spec extractions, reports)
is stored in SQLite (er_checker.db) and kept until the user removes it
from the Manage Data page.

Stack:  Python · Streamlit · SQLite · OpenRouter API (US hosts) · reportlab
Run:    streamlit run app.py
"""

import hashlib
import html
import io
import json
import os
import re
import shutil
import concurrent.futures
import sqlite3
import tempfile
import threading
import time
import unicodedata
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin, urlparse, parse_qsl, urlencode, urlunparse

import pandas as pd
import requests
import streamlit as st
from bs4 import BeautifulSoup

from components import cards, header, progress as progress_ui
from components import report as report_ui
from components import sidebar as app_sidebar
from components import styles
from components import upload as upload_ui

# ═══════════════════════════════════════════════════════════════════
#  CONSTANTS
# ═══════════════════════════════════════════════════════════════════

APP_DIR = Path(__file__).resolve().parent
DB_NAME = "er_checker.db"


def _writable_db_path() -> str:
    """Return a path where SQLite can actually write.

    On hosted platforms (e.g. Streamlit Community Cloud) the deployed
    repo folder — and any er_checker.db committed to it — can be
    read-only, which causes 'attempt to write a readonly database'.
    Try the app folder first, then fall back to the user's home dir and
    the system temp dir. If a bundled (read-only) database exists in
    the repo, seed the writable copy from it so existing data carries
    over. SQLite also needs to create journal files, so both the
    directory and the file must be writable."""
    bundled = APP_DIR / DB_NAME
    for d in (APP_DIR, Path.home() / ".er_checker",
              Path(tempfile.gettempdir()) / "er_checker"):
        try:
            d.mkdir(parents=True, exist_ok=True)
            probe = d / ".write_test"
            probe.write_text("x")
            probe.unlink()
        except OSError:
            continue  # directory not writable → journal files would fail
        p = d / DB_NAME
        if p.exists() and not os.access(p, os.W_OK):
            try:
                os.chmod(p, 0o664)          # a committed file may be mode 444
            except OSError:
                pass
        if p.exists() and not os.access(p, os.W_OK):
            continue                        # still read-only → next location
        if not p.exists() and bundled.exists() and p != bundled:
            try:
                shutil.copy(bundled, p)     # carry over bundled data
                os.chmod(p, 0o664)
            except OSError:
                pass
        return str(p)
    return str(bundled)                     # last resort


DB_PATH = _writable_db_path()


# ── remote database (Turso) ─────────────────────────────────────────
# Streamlit Community Cloud has an EPHEMERAL filesystem: every reboot /
# redeploy / wake-from-sleep starts a fresh container and deletes any
# locally written er_checker.db. To persist data there, set
#   TURSO_DATABASE_URL  (libsql://<db-name>-<org>.turso.io)
#   TURSO_AUTH_TOKEN
# in the app's Secrets. The app then uses a libsql "embedded replica":
# a local SQLite file whose writes are forwarded to the hosted Turso
# database, so data survives restarts. Without these secrets the app
# keeps using the plain local SQLite file exactly as before.

def _turso_credentials() -> tuple:
    url = token = ""
    try:
        url = str(st.secrets.get("TURSO_DATABASE_URL", "")).strip()
        token = str(st.secrets.get("TURSO_AUTH_TOKEN", "")).strip()
    except Exception:
        pass  # no secrets file — fall back to the environment
    url = url or os.environ.get("TURSO_DATABASE_URL", "").strip()
    token = token or os.environ.get("TURSO_AUTH_TOKEN", "").strip()
    return url, token


USING_TURSO = bool(_turso_credentials()[0])


class _CompatCursor:
    """Wraps a libsql cursor so rows behave like the dicts the rest of
    the app expects (row["col"], dict(row), iteration)."""

    def __init__(self, cur):
        self._cur = cur

    def _to_dict(self, tup):
        if tup is None:
            return None
        if isinstance(tup, dict):
            return tup
        cols = [d[0] for d in (self._cur.description or [])]
        return dict(zip(cols, tup))

    def fetchone(self):
        return self._to_dict(self._cur.fetchone())

    def fetchall(self):
        return [self._to_dict(t) for t in self._cur.fetchall()]

    def __iter__(self):
        while True:
            row = self.fetchone()
            if row is None:
                return
            yield row

    @property
    def lastrowid(self):
        return getattr(self._cur, "lastrowid", None)


class _TursoConn:
    """Minimal sqlite3.Connection-compatible wrapper around libsql.
    commit() also sync()s so the local replica immediately reflects the
    write that was forwarded to the hosted primary."""

    def __init__(self, conn):
        self._conn = conn

    def execute(self, sql, params=()):
        return _CompatCursor(self._conn.execute(sql, params))

    def executescript(self, script: str):
        try:
            self._conn.executescript(script)
        except AttributeError:
            for stmt in script.split(";"):
                if stmt.strip():
                    self._conn.execute(stmt)
        self.commit()

    def commit(self):
        self._conn.commit()
        try:
            self._conn.sync()
        except Exception:
            pass  # sync is best-effort; the write already reached Turso

    def close(self):
        try:
            self._conn.close()
        except Exception:
            pass


def _connect_turso():
    try:
        import libsql
    except ImportError:                       # older package name
        import libsql_experimental as libsql
    url, token = _turso_credentials()
    replica = str(Path(DB_PATH).with_name("er_checker_replica.db"))
    conn = libsql.connect(replica, sync_url=url, auth_token=token)
    try:
        conn.sync()                           # pull current data on startup
    except Exception as exc:
        print(f"Turso initial sync failed (continuing): {exc}")
    return _TursoConn(conn)


def _quarantine_bad_db(path: str, reason: str = "") -> None:
    """Move an invalid/corrupt SQLite DB aside so the app can recreate it.

    Streamlit Cloud redacts sqlite3.DatabaseError messages, but the common
    cause at startup is a damaged or non-SQLite er_checker.db file. Keeping a
    timestamped backup avoids silently deleting user data.
    """
    db = Path(path)
    if not db.exists():
        return
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    bad = db.with_name(f"{db.stem}.bad_{stamp}{db.suffix}")
    try:
        db.rename(bad)
        print(f"SQLite database was quarantined at {bad}. {reason}")
    except OSError:
        # Last resort: remove the file so init_db can create a clean DB.
        try:
            db.unlink()
            print(f"SQLite database was removed because it could not be opened. {reason}")
        except OSError:
            pass


def _ensure_sqlite_db_is_usable(path: str) -> None:
    """Validate an existing DB before Streamlit caches the connection."""
    db = Path(path)
    if not db.exists() or db.stat().st_size == 0:
        return
    try:
        conn = sqlite3.connect(path, timeout=30)
        try:
            result = conn.execute("PRAGMA quick_check").fetchone()
            if not result or str(result[0]).lower() != "ok":
                conn.close()
                _quarantine_bad_db(path, "PRAGMA quick_check failed.")
                return
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except sqlite3.DatabaseError as exc:
        _quarantine_bad_db(path, f"DatabaseError during validation: {exc}")


if not USING_TURSO:
    _ensure_sqlite_db_is_usable(DB_PATH)

USER_AGENT = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                   "Chrome/124.0 Safari/537.36")
}

# status colours (aligned with the purple theme palette)
RED = "#DC2626"
GREEN = "#16A34A"
AMBER = "#D97706"

APP_VERSION = "2.2.4"
EXTRACTION_VERSION = "2.0.3-pdf1"
ENTRY_WORDING_VERSION = "2.1.4-relevant1"

# ═══════════════════════════════════════════════════════════════════
#  DATABASE
# ═══════════════════════════════════════════════════════════════════

@st.cache_resource
def get_conn():
    if USING_TURSO:
        return _connect_turso()
    conn = sqlite3.connect(DB_PATH, timeout=30, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def _d(row):
    """sqlite3.Row → plain dict (Streamlit widgets can't deep-copy Row objects)."""
    return dict(row) if row is not None else None


def _ds(rows):
    return [dict(r) for r in rows]


@st.cache_resource(show_spinner=False)
def init_db():
    c = get_conn()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS settings (
        key   TEXT PRIMARY KEY,
        value TEXT
    );
    CREATE TABLE IF NOT EXISTS courses (
        id            INTEGER PRIMARY KEY AUTOINCREMENT,
        number        TEXT,
        name          TEXT NOT NULL,
        course_url    TEXT NOT NULL UNIQUE,
        spec_url      TEXT,
        level         TEXT,
        course_type   TEXT,
        excel_entry   TEXT,
        created_at    TEXT,
        updated_at    TEXT
    );
    CREATE TABLE IF NOT EXISTS specs (
        id            INTEGER PRIMARY KEY AUTOINCREMENT,
        url           TEXT NOT NULL UNIQUE,
        doc_hash      TEXT,
        entry_req     TEXT,
        status        TEXT DEFAULT 'pending',   -- pending | ok | error
        error         TEXT,
        extracted_at  TEXT,
        extractor_version TEXT
    );
    CREATE TABLE IF NOT EXISTS reports (
        id            INTEGER PRIMARY KEY AUTOINCREMENT,
        course_id     INTEGER NOT NULL,
        result        TEXT,       -- Pass | Fail
        page_entry    TEXT,
        spec_entry    TEXT,
        excel_entry   TEXT,
        issues_json   TEXT,
        corrected     TEXT,
        created_at    TEXT,
        FOREIGN KEY (course_id) REFERENCES courses(id) ON DELETE CASCADE
    );
    """)
    # migrations (ALTER TABLE is a no-op when the column already exists)
    migrations = (
        # cache for the formatted (point-by-point, categorised) entry reqs
        "ALTER TABLE specs ADD COLUMN entry_req_formatted TEXT",
        "ALTER TABLE specs ADD COLUMN entry_req_format_version TEXT",
        # Method of Assessment check
        "ALTER TABLE specs ADD COLUMN moa TEXT",
        "ALTER TABLE specs ADD COLUMN extractor_version TEXT",
        "ALTER TABLE reports ADD COLUMN page_moa TEXT",
        "ALTER TABLE reports ADD COLUMN spec_moa TEXT",
        "ALTER TABLE reports ADD COLUMN moa_issues_json TEXT",
        "ALTER TABLE reports ADD COLUMN moa_corrected TEXT",
        "ALTER TABLE reports ADD COLUMN moa_result TEXT",
        "ALTER TABLE reports ADD COLUMN entry_wording_version TEXT",
    )
    for m in migrations:
        try:
            c.execute(m)
        except Exception as exc:
            # Expected when the app has already run this migration before.
            # (libsql raises its own exception type, hence the broad catch.)
            if "duplicate column name" in str(exc).lower():
                pass
            else:
                raise
    c.commit()


def setting(key: str, default: str = "") -> str:
    row = get_conn().execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
    return row["value"] if row else default


def set_setting(key: str, value: str):
    c = get_conn()
    c.execute("INSERT INTO settings(key,value) VALUES(?,?) "
              "ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, value))
    c.commit()


def now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


# ── courses ─────────────────────────────────────────────────────────

def parse_level_type(name: str):
    level = course_type = ""
    m = re.search(r"level\s*(\d+)", name or "", re.I)
    if m:
        level = f"Level {m.group(1)}"
    m = re.search(r"\b(extended\s+diploma|combined\s+extended\s+diploma|diploma|"
                  r"certificate|award)\b", name or "", re.I)
    if m:
        course_type = m.group(1).title()
    return level, course_type


def upsert_course(row: dict):
    c = get_conn()
    level, ctype = row.get("level", ""), row.get("course_type", "")
    if not level or not ctype:
        pl, pt = parse_level_type(row.get("name", ""))
        level, ctype = level or pl, ctype or pt
    c.execute("""
        INSERT INTO courses (number, name, course_url, spec_url, level,
                             course_type, excel_entry, created_at, updated_at)
        VALUES (?,?,?,?,?,?,?,?,?)
        ON CONFLICT(course_url) DO UPDATE SET
            number=excluded.number, name=excluded.name,
            spec_url=excluded.spec_url, level=excluded.level,
            course_type=excluded.course_type, excel_entry=excluded.excel_entry,
            updated_at=excluded.updated_at
    """, (row.get("number", ""), row["name"], row["course_url"],
          row.get("spec_url", ""), level, ctype,
          row.get("excel_entry", ""), now(), now()))
    c.commit()


def all_courses() -> list:
    return _ds(get_conn().execute(
        "SELECT * FROM courses ORDER BY number, name").fetchall())


def get_course(cid: int):
    return _d(get_conn().execute("SELECT * FROM courses WHERE id=?", (cid,)).fetchone())


def delete_course(cid: int):
    c = get_conn()
    c.execute("DELETE FROM reports WHERE course_id=?", (cid,))
    c.execute("DELETE FROM courses WHERE id=?", (cid,))
    c.commit()


def clear_all_data():
    c = get_conn()
    for t in ("reports", "courses", "specs"):
        c.execute(f"DELETE FROM {t}")
    c.commit()


# ── specs cache ─────────────────────────────────────────────────────

def get_spec(url: str):
    return _d(get_conn().execute("SELECT * FROM specs WHERE url=?", (url,)).fetchone())


def ensure_spec_rows():
    """Make sure every distinct spec URL referenced by a course has a cache row."""
    c = get_conn()
    urls = [r["spec_url"] for r in c.execute(
        "SELECT DISTINCT spec_url FROM courses "
        "WHERE spec_url IS NOT NULL AND TRIM(spec_url) != ''")]
    for u in urls:
        c.execute("INSERT OR IGNORE INTO specs (url, status) VALUES (?, 'pending')", (u,))
    c.commit()


def all_specs() -> list:
    return _ds(get_conn().execute("SELECT * FROM specs ORDER BY id").fetchall())


def save_spec(url: str, **fields):
    c = get_conn()
    sets = ", ".join(f"{k}=?" for k in fields)
    c.execute(f"UPDATE specs SET {sets} WHERE url=?", (*fields.values(), url))
    c.commit()


def delete_spec(spec_id: int):
    c = get_conn()
    c.execute("DELETE FROM specs WHERE id=?", (spec_id,))
    c.commit()


# ── reports ─────────────────────────────────────────────────────────

def save_report(course_id: int, result: str, page_entry: str, spec_entry: str,
                excel_entry: str, issues: dict, corrected: str,
                page_moa: str = "", spec_moa: str = "", moa_issues: dict = None,
                moa_corrected: str = "", moa_result: str = "",
                entry_wording_version: str = ENTRY_WORDING_VERSION) -> int:
    c = get_conn()
    cur = c.execute("""
        INSERT INTO reports (course_id, result, page_entry, spec_entry,
                             excel_entry, issues_json, corrected, created_at,
                             page_moa, spec_moa, moa_issues_json,
                             moa_corrected, moa_result, entry_wording_version)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (course_id, result, page_entry, spec_entry, excel_entry,
          json.dumps(issues, ensure_ascii=False), corrected, now(),
          page_moa, spec_moa,
          json.dumps(moa_issues or {}, ensure_ascii=False),
          moa_corrected, moa_result, entry_wording_version))
    c.commit()
    return cur.lastrowid


def latest_report(course_id: int):
    return _d(get_conn().execute(
        "SELECT * FROM reports WHERE course_id=? ORDER BY id DESC LIMIT 1",
        (course_id,)).fetchone())


def all_reports() -> list:
    return _ds(get_conn().execute("""
        SELECT r.*, c.name AS course_name, c.course_url, c.spec_url
        FROM reports r JOIN courses c ON c.id = r.course_id
        ORDER BY r.id DESC
    """).fetchall())


def delete_report(report_id: int):
    c = get_conn()
    c.execute("DELETE FROM reports WHERE id=?", (report_id,))
    c.commit()


# ═══════════════════════════════════════════════════════════════════
#  EXTRACTION — spec documents
# ═══════════════════════════════════════════════════════════════════

def _drive_file_id(url: str):
    """Extract the file id from any common Google Drive link shape."""
    for pat in (r"drive\.google\.com/file/d/([^/?#]+)",
                r"drive\.google\.com/(?:open|uc)\?[^#]*?id=([^&#]+)",
                r"drive\.usercontent\.google\.com/download\?[^#]*?id=([^&#]+)"):
        m = re.search(pat, url or "")
        if m:
            return m.group(1)
    return None


def _normalise_document_bytes(data: bytes) -> bytes:
    """Trim harmless bytes before a PDF header.

    Some document servers prepend a UTF-8 BOM, whitespace, or a short proxy
    banner before ``%PDF``. Browsers tolerate this, but strict type checks do
    not. Only trim when the real header appears near the beginning.
    """
    if not data:
        return data
    idx = data[:4096].find(b"%PDF-")
    if idx > 0:
        return data[idx:]
    return data


def _looks_like_html(data: bytes, content_type: str = "") -> bool:
    head = (data or b"")[:8000].lstrip().lower()
    ctype = (content_type or "").lower()
    return ("text/html" in ctype or "application/xhtml" in ctype
            or head.startswith(b"<!doctype html") or head.startswith(b"<html")
            or b"<html" in head[:1000])


def _html_access_problem(data: bytes):
    """Explain common cases where a browser page is returned, not a file."""
    text = BeautifulSoup((data or b"").decode("utf-8", errors="ignore"),
                         "html.parser").get_text(" ", strip=True).lower()
    sample = text[:12000]
    if any(x in sample for x in (
            "sign in to continue", "please sign in", "log in to continue",
            "authentication required", "you must be logged in")):
        return ("The link opens only for a signed-in browser. The app has no "
                "access to your browser cookies; publish the document for "
                "anonymous access or use a direct public download link.")
    if any(x in sample for x in (
            "access denied", "request access", "you do not have permission",
            "you need permission", "unauthorized", "forbidden")):
        return ("The document server denied anonymous access. Set the file to "
                "public/'Anyone with the link' or provide a direct public link.")
    if any(x in sample for x in (
            "checking your browser", "verify you are human", "captcha",
            "enable javascript and cookies", "cloudflare")):
        return ("The site returned an anti-bot/JavaScript challenge instead of "
                "the document. A direct public PDF link is required.")
    return None


def _embedded_document_url(data: bytes, base_url: str):
    """Find a PDF/DOCX embedded in a browser viewer or landing page."""
    try:
        soup = BeautifulSoup(data.decode("utf-8", errors="ignore"), "html.parser")
    except Exception:
        return None

    raw = []
    for tag, attr in (("embed", "src"), ("object", "data"),
                      ("iframe", "src"), ("a", "href")):
        for node in soup.find_all(tag):
            value = node.get(attr)
            if value:
                raw.append(value)
    for node in soup.find_all("meta"):
        value = node.get("content")
        if value:
            raw.append(value)

    # Prefer explicit document-looking URLs. Ignore javascript/data URLs and
    # the original landing page to prevent loops.
    base_no_frag = (base_url or "").split("#", 1)[0]
    candidates = []
    for value in raw:
        value = str(value).strip().replace("&amp;", "&")
        if not value or value.lower().startswith(("javascript:", "data:")):
            continue
        absolute = urljoin(base_url, value).split("#", 1)[0]
        if absolute == base_no_frag:
            continue
        low = absolute.lower()
        score = 0
        if re.search(r"\.(pdf|docx?)(?:$|[?#])", low):
            score += 10
        if any(k in low for k in ("download=1", "download=true", "export=download",
                                  "format=pdf", "/download")):
            score += 5
        if score:
            candidates.append((score, absolute))
    return max(candidates, default=(0, None))[1]


def _with_query_param(url: str, key: str, value: str) -> str:
    parts = urlparse(url)
    query = dict(parse_qsl(parts.query, keep_blank_values=True))
    query[key] = value
    return urlunparse(parts._replace(query=urlencode(query)))


def _ordinary_download_candidates(url: str) -> list[str]:
    """Return common direct-download variants before the original URL."""
    out = []
    host = urlparse(url or "").netloc.lower()
    if "dropbox.com" in host:
        out.append(_with_query_param(url, "dl", "1"))
    if any(x in host for x in ("sharepoint.com", "1drv.ms", "onedrive.live.com")):
        out.append(_with_query_param(url, "download", "1"))
    out.append(url)
    # Preserve order while removing duplicates.
    return list(dict.fromkeys(x for x in out if x))


def _get_document_response(url: str, *, referer: str = ""):
    """GET a document with limited retry and clearer HTTP errors."""
    headers = dict(USER_AGENT)
    headers.update({
        "Accept": ("application/pdf,application/vnd.openxmlformats-officedocument."
                   "wordprocessingml.document,text/html;q=0.8,*/*;q=0.5"),
        "Accept-Language": "en-GB,en;q=0.9",
    })
    if referer:
        headers["Referer"] = referer

    last = None
    for attempt in range(3):
        try:
            resp = requests.get(url, headers=headers, timeout=(20, 90),
                                allow_redirects=True)
            if resp.status_code in (429, 500, 502, 503, 504) and attempt < 2:
                time.sleep(0.8 * (attempt + 1))
                last = resp
                continue
            if resp.status_code in (401, 403):
                raise RuntimeError(
                    f"Document server returned HTTP {resp.status_code}. The link "
                    "may work in your browser because you are signed in, but the "
                    "app needs an anonymously accessible direct download link.")
            if resp.status_code == 404:
                raise RuntimeError("Document server returned HTTP 404 (file not found).")
            resp.raise_for_status()
            return resp
        except requests.RequestException as exc:
            last = exc
            if attempt < 2:
                time.sleep(0.8 * (attempt + 1))
                continue
            raise RuntimeError(f"Could not download the document: {exc}") from exc
    if hasattr(last, "raise_for_status"):
        last.raise_for_status()
    raise RuntimeError("Could not download the document.")


def _drive_html_problem(data: bytes):
    """If Drive returned an HTML page instead of the file, say why."""
    head = data[:12000]
    if not _looks_like_html(head):
        return None
    low = head.lower()
    if b"virus scan warning" in low or b"confirm=" in low:
        return "scan-page"          # bypassable — handled by caller
    generic = _html_access_problem(data)
    if generic:
        return generic
    if b"too many users have viewed or downloaded" in low or b"quota" in low:
        return ("Google Drive download quota exceeded for this file — try "
                "again later or re-upload a copy.")
    if b"file does not exist" in low or b"not found" in low:
        return "Google Drive file not found — the link may be broken."
    return ("Google Drive returned a viewer page instead of the document. "
            "Set sharing to 'Anyone with the link' and use a file share link.")


def fetch_spec_bytes(url: str) -> bytes:
    """Download a PDF, DOCX, or specification web page reliably.

    Supports Google Drive/Docs, common cloud-share direct-download variants,
    and HTML viewer pages that contain an embedded PDF link.
    """
    if not (url or "").strip():
        raise RuntimeError("No specification document URL was provided.")

    # Google Docs native document → export as PDF
    m = re.search(r"docs\.google\.com/document/d/([^/?#]+)", url or "")
    if m:
        u = f"https://docs.google.com/document/d/{m.group(1)}/export?format=pdf"
        resp = _get_document_response(u, referer=url)
        data = _normalise_document_bytes(resp.content)
        issue = _drive_html_problem(data)
        if issue:
            raise RuntimeError(issue)
        return data

    fid = _drive_file_id(url)
    if fid:
        candidates = [
            f"https://drive.usercontent.google.com/download?id={fid}"
            f"&export=download&confirm=t",
            f"https://drive.google.com/uc?export=download&id={fid}&confirm=t",
        ]
        problem = None
        for u in candidates:
            try:
                resp = _get_document_response(u, referer=url)
            except Exception as exc:
                problem = str(exc)
                continue
            data = _normalise_document_bytes(resp.content)
            issue = _drive_html_problem(data)
            if issue is None:
                return data
            if issue == "scan-page":
                mm = re.search(rb'href="([^"]*confirm=[^"]*)"', data)
                if mm:
                    cu = mm.group(1).decode().replace("&amp;", "&")
                    if cu.startswith("/"):
                        cu = "https://drive.google.com" + cu
                    try:
                        r2 = _get_document_response(cu, referer=url)
                        data2 = _normalise_document_bytes(r2.content)
                        if _drive_html_problem(data2) is None:
                            return data2
                    except Exception as exc:
                        problem = str(exc)
                problem = problem or "Could not pass the Google Drive scan page."
            else:
                problem = issue
        raise RuntimeError(problem or "Google Drive download failed.")

    # Ordinary links. Try cloud-provider direct download variants first.
    problems = []
    for candidate in _ordinary_download_candidates(url):
        try:
            resp = _get_document_response(candidate, referer=url if candidate != url else "")
        except Exception as exc:
            problems.append(str(exc))
            continue
        data = _normalise_document_bytes(resp.content)
        ctype = resp.headers.get("Content-Type", "")

        if data.startswith((b"%PDF-", b"PK")):
            return data

        if _looks_like_html(data, ctype):
            access_problem = _html_access_problem(data)
            nested = _embedded_document_url(data, resp.url or candidate)
            if nested:
                try:
                    nested_resp = _get_document_response(nested,
                                                         referer=resp.url or candidate)
                    nested_data = _normalise_document_bytes(nested_resp.content)
                    if nested_data.startswith((b"%PDF-", b"PK")):
                        return nested_data
                    if not _looks_like_html(
                            nested_data, nested_resp.headers.get("Content-Type", "")):
                        return nested_data
                except Exception as exc:
                    problems.append(f"Embedded document download failed: {exc}")
            if access_problem:
                problems.append(access_problem)
                continue
            # A genuine specification web page is valid input.
            return data

        # Plain text and unusual but readable formats are passed through to the
        # text converter, which provides the final type/error diagnosis.
        return data

    raise RuntimeError(problems[-1] if problems else "Document download failed.")


def _extract_pdf_with_pdfplumber(data: bytes, max_chars: int) -> tuple[str, int]:
    import pdfplumber
    parts = []
    pages = 0
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = len(pdf.pages)
        for page in pdf.pages:
            parts.append(page.extract_text(x_tolerance=2, y_tolerance=3) or "")
            if sum(len(p) for p in parts) > max_chars:
                break
    return "\n".join(parts)[:max_chars], pages


def _extract_pdf_with_pypdf(data: bytes, max_chars: int) -> tuple[str, int]:
    """Fallback for PDFs that pdfplumber/pdfminer cannot parse well."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data), strict=False)
    if reader.is_encrypted:
        try:
            result = reader.decrypt("")
        except Exception as exc:
            raise RuntimeError(
                "The PDF is password-protected and cannot be extracted.") from exc
        if result == 0:
            raise RuntimeError(
                "The PDF is password-protected and cannot be extracted.")
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
        if sum(len(p) for p in parts) > max_chars:
            break
    return "\n".join(parts)[:max_chars], len(reader.pages)


def spec_bytes_to_text(data: bytes, url: str = "", max_chars: int = 150000) -> str:
    """Convert PDF / DOCX / HTML / text bytes into plain text.

    PDF extraction uses two independent engines. A clear error is returned for
    scanned/image-only or protected PDFs instead of silently caching blank text.
    """
    data = _normalise_document_bytes(data)
    if not data:
        raise RuntimeError("The downloaded document is empty.")

    if data.startswith(b"%PDF-"):
        attempts = []
        best_text = ""
        page_count = 0
        try:
            text, pages = _extract_pdf_with_pdfplumber(data, max_chars)
            best_text, page_count = text, max(page_count, pages)
        except Exception as exc:
            attempts.append(f"pdfplumber: {exc}")
        try:
            text, pages = _extract_pdf_with_pypdf(data, max_chars)
            page_count = max(page_count, pages)
            if len(re.sub(r"\s+", "", text)) > len(re.sub(r"\s+", "", best_text)):
                best_text = text
        except Exception as exc:
            attempts.append(f"pypdf: {exc}")

        clean_len = len(re.sub(r"\s+", "", best_text))
        if clean_len < 30:
            protected = next((x for x in attempts if "password-protected" in x), None)
            if protected:
                raise RuntimeError("The PDF is password-protected and cannot be extracted.")
            if page_count:
                raise RuntimeError(
                    f"The PDF downloaded and opened ({page_count} page(s)), but it "
                    "contains no usable machine-readable text. It is probably a "
                    "scanned/image-only PDF. Convert it to a searchable/OCR PDF "
                    "or upload a text-based copy.")
            detail = "; ".join(attempts)[:500]
            raise RuntimeError(
                "The file looks like a PDF but could not be parsed. "
                + (f"Parser details: {detail}" if detail else
                   "It may be damaged or use an unsupported PDF structure."))
        return best_text[:max_chars]

    if data.startswith(b"PK"):
        from docx import Document
        try:
            d = Document(io.BytesIO(data))
        except Exception as exc:
            raise RuntimeError(
                "The downloaded ZIP-based file is not a readable DOCX document.") from exc
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(p for p in parts if p and p.strip())[:max_chars]
        if not text.strip():
            raise RuntimeError("The DOCX opened but contained no extractable text.")
        return text

    if _looks_like_html(data):
        problem = _html_access_problem(data)
        if problem:
            raise RuntimeError(problem)
        soup = BeautifulSoup(data.decode("utf-8", errors="ignore"), "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
            tag.decompose()
        text = soup.get_text("\n", strip=True)[:max_chars]
        if len(re.sub(r"\s+", "", text)) < 30:
            raise RuntimeError(
                "The link returned a web viewer/landing page with no extractable "
                "document text. Use the viewer's direct Download link.")
        return text

    # Last-resort plain-text decoding for servers that omit or mislabel MIME.
    text = data.decode("utf-8", errors="ignore")[:max_chars]
    if len(re.sub(r"\s+", "", text)) < 30:
        raise RuntimeError(
            "The downloaded file is not a recognised text PDF, DOCX, or HTML document.")
    return text


ENTRY_HEADING = r"entry\s+requirements?|entry\s+criteria|admission\s+requirements?"
MOA_HEADING = (r"method(?:s)?\s+of\s+assessment"
               r"|assessment\s+(?:methods?|approach|overview|strategy|and\s+grading)"
               r"|how\s+(?:is\s+(?:this|the)\s+course|will\s+I\s+be)\s+assessed")
OVERVIEW_HEADING = (
    r"course\s+overview|overview|course\s+description|about\s+(?:this|the)\s+course"
    r"|what\s+is\s+this\s+course\s+about|introduction"
)

# Words/phrases that begin the NEXT section of a spec document.
# NOTE: bare generic words (resources, support, contact, assessment, units…)
# previously matched ordinary body-text lines — PDF extraction preserves the
# document's visual line-wrapping, so a sentence like
#   "…to enable them to access relevant\nresources and complete the unit assignments."
# was mistaken for a "Resources" heading and truncated the section, silently
# dropping everything after it (e.g. the IELTS / CEFR / CAE / PTE list).
_COMMON_STOPS = (
    r"progression(?:\s+(?:opportunities|routes?))?|grading"
    r"|units?(?:\s+(?:overview|summary))?|qualification\s+structure|unit\s+structure"
    r"|guided\s+learning(?:\s+hours)?|total\s+qualification\s+time"
    r"|course\s+content|learning\s+outcomes?|funding|centre(?:\s+requirements?)?"
    r"|appendix(?:\s+\w+)?|introduction|support(?:\s+for\s+learners)?"
    r"|resources?(?:\s+required)?|contact(?:\s+us)?"
    r"|reasonable\s+adjustments|recognition\s+of\s+prior(?:\s+learning)?"
)
# next-section words when extracting ENTRY REQUIREMENTS (assessment ends it)
_STOP_WORDS = (
    r"method(?:s)?\s+of\s+assessment"
    r"|assessment(?:\s+(?:overview|methods?|and\s+grading|strategy))?|"
    + _COMMON_STOPS
)
# next-section words when extracting METHOD OF ASSESSMENT (entry reqs end it)
_MOA_STOP_WORDS = (
    rf"{ENTRY_HEADING}|malpractice|certification|results|external\s+verification|"
    + _COMMON_STOPS
)


def _stop_re(words: str) -> str:
    # A stop match only counts as a heading when it occupies (almost) the
    # whole line: optional section number before, optional colon after, then
    # end of line. A wrapped sentence continuing after the word will NOT match.
    return (rf"(?:^|\n)[ \t]*(?:\d+(?:\.\d+)*\.?[ \t]+)?"
            rf"(?:{words})"
            rf"[ \t]*:?[ \t]*(?=\n|$)")


STOP_HEADINGS = _stop_re(_STOP_WORDS)
MOA_STOP_HEADINGS = _stop_re(_MOA_STOP_WORDS)


def heuristic_section(text: str, heading_re: str, stop_re: str,
                      max_chars: int = 6000) -> str:
    """Pull a named section out of a document's text.
    Skips table-of-contents hits (dot leaders / bare page numbers) and,
    when the heading appears more than once, keeps the longest candidate
    (body section) rather than the first (often a passing mention)."""
    if not text:
        return ""
    best = ""
    for m in re.finditer(rf"(?:^|\n)\s*(?:\d+(?:\.\d+)*\.?\s+)?(?:{heading_re})"
                         rf"\s*:?\s*\n?", text, re.I):
        start = m.end()
        stop = re.search(stop_re, text[start:], re.I)
        chunk = text[start:start + stop.start()] if stop else text[start:start + max_chars]
        chunk = re.sub(r"\n{3,}", "\n\n", chunk).strip()[:max_chars]
        head = chunk[:80]
        # TOC lines look like "..... 7" or just a page number
        if re.match(r"^[.·\s]*\d{1,3}\s*$", head) or re.match(r"^\.{4,}", head):
            continue
        if len(chunk) > len(best):
            best = chunk
    return best if len(best) >= 40 else ""


def heuristic_entry_section(text: str, max_chars: int = 6000) -> str:
    return heuristic_section(text, ENTRY_HEADING, STOP_HEADINGS, max_chars)


def heuristic_moa_section(text: str, max_chars: int = 6000) -> str:
    return heuristic_section(text, MOA_HEADING, MOA_STOP_HEADINGS, max_chars)


# ═══════════════════════════════════════════════════════════════════
#  EXTRACTION — course pages
# ═══════════════════════════════════════════════════════════════════

# headings that mark the NEXT section on a course page (used to stop reading)
_PAGE_NEXT = (r"course (?:content|curriculum)|qualification|awarding body"
              r"|career path|who is this|certification|progression|why study"
              r"|faq|average completion|not sure if this course"
              r"|speak to an advisor|study method|course duration")
PAGE_NEXT_AFTER_ENTRY = rf"method(?:s)? of assessment|assessment|{_PAGE_NEXT}"
PAGE_NEXT_AFTER_MOA = rf"entry requirements?|entry criteria|{_PAGE_NEXT}"
PAGE_NEXT_AFTER_OVERVIEW = (
    rf"{ENTRY_HEADING}|{MOA_HEADING}|course\s+(?:content|curriculum|modules?)"
    rf"|who\s+is\s+this\s+(?:course\s+)?for|career(?:s|\s+path)?|progression"
    rf"|qualification|awarding\s+body|study\s+method|course\s+duration|faq"
)

# trailing marketing boilerplate to trim from any extracted page section
_PAGE_BOILERPLATE = (r"not sure if this course|speak to an advisor"
                     r"|average completion timeframe|\b0\d{2}[- ]\d{4}[- ]\d{4}\b")


def _normalise_heading_text(text: str) -> str:
    """Collapse whitespace so heading comparisons are exact and predictable."""
    return re.sub(r"\s+", " ", text or "").strip()


def _matches_heading(text: str, pattern: str) -> bool:
    text = _normalise_heading_text(text)
    return bool(text and re.fullmatch(rf"(?:{pattern})\s*:?\s*", text, re.I))


def _is_heading_node(node, pattern: str) -> bool:
    """Return True only when *node itself* looks like a section heading.

    The old extractor used ``re.match`` against every short sibling. That made
    ordinary requirement lines such as "Qualifications at Level 3 ..." look
    like a new "Qualification" section and silently truncated the result.
    """
    if node is None or not getattr(node, "name", None):
        return False
    text = _normalise_heading_text(node.get_text(" ", strip=True))
    if not _matches_heading(text, pattern):
        return False
    if node.name in {"h1", "h2", "h3", "h4", "h5", "h6", "button", "summary"}:
        return True
    role = str(node.attrs.get("role", "")).lower()
    classes = " ".join(node.attrs.get("class", [])).lower()
    if role in {"heading", "button", "tab"}:
        return True
    if any(word in classes for word in ("heading", "title", "accordion", "tab")):
        return True
    # Some CMSs use a plain <strong>/<b>/<p>/<div> as a heading. Exact full
    # matching keeps this safe; body lines with extra wording cannot match.
    return node.name in {"strong", "b", "p", "div", "span", "a"} and len(text) < 90


def _clean_section_candidate(text: str, heading_re: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text or "").strip()
    text = re.sub(rf"^\s*(?:{heading_re})\s*:?\s*", "", text, flags=re.I)
    return text.strip()[:4000]


def _accordion_target(soup, heading) -> str:
    """Read content referenced by an accordion/tab heading when available."""
    ids = []
    for attr in ("aria-controls", "data-target", "data-bs-target", "href"):
        value = heading.attrs.get(attr)
        if isinstance(value, str) and value.startswith("#"):
            ids.append(value[1:])
        elif attr == "aria-controls" and isinstance(value, str):
            ids.append(value)
    for target_id in ids:
        target = soup.find(id=target_id)
        if target is not None:
            return target.get_text("\n", strip=True)
    return ""


def _page_section(soup, heading_re: str, next_re: str) -> str:
    """Extract the most complete matching course-page section.

    Supports ordinary headings, nested cards, and accordion/tab controls. All
    matching headings are considered and the best body candidate is selected,
    avoiding menu/summary duplicates that appear before the real content.
    """
    candidates = []
    heading_tags = ["h1", "h2", "h3", "h4", "h5", "h6", "strong", "b",
                    "button", "summary", "a", "span", "p", "div"]
    headings = []
    for h in soup.find_all(heading_tags):
        text = _normalise_heading_text(h.get_text(" ", strip=True))
        if len(text) < 90 and _matches_heading(text, heading_re):
            headings.append(h)

    for heading in headings:
        # 1) Explicit accordion / tab target.
        target_text = _accordion_target(soup, heading)
        if target_text:
            candidates.append(_clean_section_candidate(target_text, heading_re))

        # 2) Siblings at the heading level. Stop only at a genuine heading node
        # whose complete text matches the next-section pattern.
        parts = []
        node = heading
        for _ in range(35):
            node = node.find_next_sibling()
            if node is None:
                break
            if _is_heading_node(node, next_re):
                break
            text = node.get_text("\n", strip=True)
            if text:
                parts.append(text)
            if sum(len(part) for part in parts) > 4000:
                break
        if parts:
            candidates.append(_clean_section_candidate("\n".join(parts), heading_re))

        # 3) Nested card/accordion layouts where heading and content share a
        # parent. Try a few ancestor levels, but reject containers that include
        # another recognised section heading after the current one.
        parent = heading.parent
        for _ in range(3):
            if parent is None or getattr(parent, "name", None) in {"body", "main", "article"}:
                break
            parent_text = parent.get_text("\n", strip=True)
            cleaned = _clean_section_candidate(parent_text, heading_re)
            if cleaned and len(cleaned) > 20:
                # Trim at an exact next heading line if the container holds
                # multiple sections.
                lines = cleaned.splitlines()
                kept = []
                for line in lines:
                    if _matches_heading(line, next_re):
                        break
                    kept.append(line)
                candidates.append("\n".join(kept).strip()[:4000])
            parent = parent.parent

    trimmed_candidates = []
    for candidate in candidates:
        kept = []
        for line in candidate.splitlines():
            if _matches_heading(line, next_re):
                break
            kept.append(line)
        value = "\n".join(kept).strip()[:4000]
        if value and len(value) >= 15:
            trimmed_candidates.append(value)
    candidates = trimmed_candidates
    if not candidates:
        return ""

    # Prefer informative, complete candidates without rewarding duplicated
    # navigation text. Length is the main signal, with line diversity as a tie
    # breaker.
    def score(value: str):
        lines = [ln.strip() for ln in value.splitlines() if ln.strip()]
        unique = len(dict.fromkeys(lines))
        return (min(len(value), 4000), unique)

    return max(candidates, key=score)


def _trim_boilerplate(text: str) -> str:
    return re.split(_PAGE_BOILERPLATE, text or "", flags=re.I)[0].strip()[:4000]


def extract_page_sections(url: str) -> tuple:
    """Return (entry_requirements, method_of_assessment, full_page_text)
    from a course page, fetched once. Heading-based extraction handles the
    usual layouts; when it fails, a heuristic runs on the flattened text,
    and the caller can still fall back to sending full_page_text to the AI."""
    try:
        resp = requests.get(url, headers=USER_AGENT, timeout=60,
                            allow_redirects=True)
        resp.raise_for_status()
    except requests.RequestException as exc:
        raise RuntimeError(
            f"Could not open the course page URL: {url}. "
            f"The site may be blocking automated access or the URL may be "
            f"unavailable. Details: {exc}"
        ) from exc
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "form"]):
        tag.decompose()

    entry = _page_section(soup, ENTRY_HEADING, PAGE_NEXT_AFTER_ENTRY)
    moa = _page_section(soup, MOA_HEADING, PAGE_NEXT_AFTER_MOA)

    main = soup.find("main") or soup.find("article") or soup.body or soup
    # Retain a wider text window. The AI helper below focuses around the
    # relevant heading before sending at most 14k characters to the model.
    full_text = re.sub(r"\n{3,}", "\n\n", main.get_text("\n", strip=True))[:40000]

    if not entry:
        entry = heuristic_entry_section(full_text)
    if not moa:
        moa = heuristic_moa_section(full_text)

    return _trim_boilerplate(entry), _trim_boilerplate(moa), full_text


def section_ai_window(text: str, heading_re: str, max_chars: int = 14000) -> str:
    """Focus a long page around a requested section before AI extraction."""
    text = text or ""
    matches = list(re.finditer(heading_re, text, re.I))
    if not matches:
        return text[:max_chars]
    # Navigation/header/footer elements were already removed from the page;
    # when a CMS repeats the label inside the main content, the later match is
    # usually the expanded section rather than a summary link.
    match = matches[-1]
    start = max(0, match.start() - 800)
    return text[start:start + max_chars]


def reconcile_section(primary: str, ai_value: str) -> str:
    """Choose the more complete extraction without replacing good DOM text
    with a tiny or obviously over-broad AI answer."""
    primary = _trim_boilerplate(primary)
    ai_value = _trim_boilerplate(ai_value)
    if not primary:
        return ai_value
    if not ai_value:
        return primary
    if len(primary) < 40 and len(ai_value) > len(primary):
        return ai_value
    # Require a meaningful completeness gain. Cap protects against an AI reply
    # that accidentally includes the rest of the page.
    if len(ai_value) <= 5000 and len(ai_value) >= len(primary) + max(60, int(len(primary) * 0.15)):
        return ai_value
    return primary


def extract_page_entry(url: str) -> tuple:
    """Backward-compatible wrapper: (entry_requirements_text, full_page_text)."""
    entry, _, full_text = extract_page_sections(url)
    return entry, full_text


def extract_page_overview(url: str) -> tuple:
    """Return (overview_text, full_page_text) from a public course webpage.

    The extractor first looks for common overview/description headings. If the
    page has no explicit overview heading, it uses the introductory text under
    the main H1 and the page meta description as fallbacks.
    """
    try:
        resp = requests.get(url, headers=USER_AGENT, timeout=60,
                            allow_redirects=True)
        resp.raise_for_status()
    except requests.RequestException as exc:
        raise RuntimeError(
            f"Could not open the selected course page: {url}. Details: {exc}"
        ) from exc

    soup = BeautifulSoup(resp.text, "html.parser")
    meta_description = ""
    meta = soup.find("meta", attrs={"name": re.compile(r"^description$", re.I)})
    if meta and meta.get("content"):
        meta_description = _normalise_heading_text(meta.get("content"))

    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "form"]):
        tag.decompose()

    overview = _page_section(soup, OVERVIEW_HEADING, PAGE_NEXT_AFTER_OVERVIEW)
    main = soup.find("main") or soup.find("article") or soup.body or soup
    full_text = re.sub(r"\n{3,}", "\n\n", main.get_text("\n", strip=True))[:40000]

    if not overview:
        h1 = main.find("h1") if hasattr(main, "find") else None
        parts = []
        if h1:
            node = h1
            for _ in range(18):
                node = node.find_next_sibling()
                if node is None:
                    break
                if getattr(node, "name", None) in {"h2", "h3", "h4"}:
                    break
                value = node.get_text("\n", strip=True)
                if value:
                    parts.append(value)
                if sum(len(x) for x in parts) >= 3500:
                    break
        intro = "\n".join(parts).strip()
        if len(intro) >= 80:
            overview = intro
        elif len(meta_description) >= 40:
            overview = meta_description

    overview = _trim_boilerplate(overview)
    return overview[:6000], full_text


# ═══════════════════════════════════════════════════════════════════
#  AI (OpenRouter · US providers)
# ═══════════════════════════════════════════════════════════════════

# ── LLM configuration — OpenRouter, pinned to US-hosted providers ────
# The provider is fixed internally; there is no provider or API-key UI.
# Requests go through OpenRouter but are ONLY served by US hosts
# (DeepInfra first, then Fireworks/Together) — never DeepSeek's own API.
# The API key is read from Streamlit Secrets (.streamlit/secrets.toml):
#   OPENROUTER_API_KEY = "sk-or-..."
# (falls back to the OPENROUTER_API_KEY environment variable for local dev)

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL = "deepseek/deepseek-v4-flash"
US_PROVIDERS = ["deepinfra", "fireworks", "together"]  # preferred order


def openrouter_api_key() -> str:
    try:
        if "OPENROUTER_API_KEY" in st.secrets:
            return str(st.secrets["OPENROUTER_API_KEY"]).strip()
    except Exception:
        pass  # no secrets file — fall back to the environment
    return os.environ.get("OPENROUTER_API_KEY", "").strip()


# gentle client-side pacing + automatic retry, because each check makes
# several AI calls (now in parallel) and OpenRouter rate-limits by tier
_MIN_CALL_GAP = 0.35     # seconds between request starts (thread-safe)
_MAX_RETRIES = 5         # attempts on 429 / 5xx before giving up
_last_call_ts = [0.0]
_call_lock = threading.Lock()


def call_ai(prompt: str, system: str, temperature: float = 0.0) -> str:
    api_key = openrouter_api_key()
    if not api_key:
        raise RuntimeError("No OpenRouter API key found — add OPENROUTER_API_KEY "
                           "to Streamlit Secrets.")
    payload = {
        "model": OPENROUTER_MODEL,
        "temperature": temperature,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": prompt},
        ],
        # pin to US hosts: DeepInfra preferred; never any other provider
        "provider": {
            "order": US_PROVIDERS,
            "only": US_PROVIDERS,
            "allow_fallbacks": False,
        },
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://streamlit.io",
        "X-Title": "Course Content Checker",
    }

    last_resp = None
    for attempt in range(_MAX_RETRIES):
        # pace request starts so parallel bursts don't trip the rate limit
        with _call_lock:
            gap = _MIN_CALL_GAP - (time.time() - _last_call_ts[0])
            if gap > 0:
                time.sleep(gap)
            _last_call_ts[0] = time.time()
        resp = requests.post(OPENROUTER_URL, headers=headers, json=payload,
                             timeout=180)
        last_resp = resp
        if resp.status_code == 429 or resp.status_code >= 500:
            # honour Retry-After when present, else exponential backoff
            try:
                wait = float(resp.headers.get("Retry-After", ""))
            except (TypeError, ValueError):
                wait = 2.0 * (2 ** attempt)          # 2, 4, 8, 16, 32 s
            time.sleep(min(wait, 45))
            continue
        if resp.status_code >= 400:
            try:
                payload_error = resp.json().get("error", {})
                message = payload_error.get("message") or str(payload_error)
            except Exception:
                message = (getattr(resp, "text", "") or "").strip()
            message = message or "No error details were returned."
            hints = {
                401: "Check that OPENROUTER_API_KEY is valid and active.",
                402: "The OpenRouter account has insufficient credits.",
                403: "The API key does not have permission for this request.",
                404: "The configured model or endpoint was not found.",
            }
            hint = hints.get(resp.status_code, "")
            raise RuntimeError(
                f"OpenRouter request failed ({resp.status_code}): {message}"
                + (f" {hint}" if hint else "")
            )
        data = resp.json()
        try:
            return data["choices"][0]["message"]["content"]
        except (KeyError, IndexError, TypeError) as exc:
            raise RuntimeError(
                "OpenRouter returned an unexpected response without model text."
            ) from exc

    # retries exhausted — raise a helpful error
    if last_resp is not None and last_resp.status_code == 429:
        raise RuntimeError(
            "OpenRouter rate limit (429) persisted after retries. Wait a "
            "minute and try again — or add credits to your OpenRouter "
            "account to raise the per-minute/day limits.")
    last_resp.raise_for_status()
    return ""


def parse_json_reply(raw: str) -> dict:
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip(), flags=re.S)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", raw, flags=re.S)
        if m:
            return json.loads(m.group(0))
        raise


AI_EXTRACT_SYSTEM = ("You extract sections from documents. "
                     "Reply ONLY with valid JSON — no markdown, no commentary.")

AI_EXTRACT_PROMPT = """From the text below, extract the {section} section verbatim (fix broken line-wrapping but do not reword). If it is genuinely absent, use an empty string.

Reply with EXACTLY this JSON: {{"{key}": "..."}}

TEXT:
{text}
"""


def ai_extract_section(text: str, section: str, key: str) -> str:
    raw = call_ai(AI_EXTRACT_PROMPT.format(section=section, key=key,
                                           text=text[:14000]),
                  AI_EXTRACT_SYSTEM)
    return parse_json_reply(raw).get(key, "").strip()


def ai_extract_entry(text: str) -> str:
    return ai_extract_section(text, "ENTRY REQUIREMENTS", "entry_requirements")


def ai_extract_moa(text: str) -> str:
    return ai_extract_section(text, "METHOD OF ASSESSMENT (how learners are "
                              "assessed: assignments, exams, coursework, "
                              "internal/external assessment, grading approach)",
                              "method_of_assessment")


AI_COMPARE_SYSTEM = (
    "You are a meticulous quality auditor for a UK college. You compare a course "
    "page's Entry Requirements against the awarding body's Qualification "
    "Specification and the internal tracker sheet. Reply ONLY with valid JSON — "
    "no markdown fences, no commentary."
)

AI_COMPARE_PROMPT = """Course: "{name}"

Compare the Entry Requirements from the three sources below. The QUALIFICATION SPECIFICATION is the authoritative source of truth; the COURSE PAGE is what students see and must faithfully reflect the specification; the EXCEL TRACKER is the internal record.

Identify:
1. Whether the course page's entry requirements match the specification (and whether the Excel tracker matches too).
2. Missing requirements — present in the specification but absent from the course page.
3. Incorrect requirements — present on the course page but wrong or absent per the specification (e.g. wrong age, wrong grades, wrong test scores).
4. Wording differences that change the meaning (e.g. "must" vs "recommended", "and" vs "or").
5. Grammar and spelling issues on the course page.
Decision rules:
* The extracted specification section may contain neighbouring or administrative text. Treat ONLY explicit applicant-facing admission or eligibility conditions as Entry Requirements.
* Include criteria such as minimum age, prerequisite qualifications or grades, English/literacy/numeracy standards, required experience, interviews, portfolios, auditions, references, checks, and explicitly accepted alternative entry routes.
* Ignore course content, units, assessment, grading, duration, progression, careers, marketing text, general learner-profile descriptions, centre/provider responsibilities, quality-assurance processes, registration administration, resources, reasonable adjustments, and special considerations unless a sentence explicitly states an applicant admission condition.
* When the qualification specification is available, it is the ONLY authority for the course-page Pass/Fail result. An Excel mismatch must be reported through matches_excel/summary but must not by itself fail the course page.
* When the qualification specification is unavailable, use the Excel tracker as the fallback authority. State this clearly in the summary.
* Compare only explicit requirements. Do not invent, infer, or generalise requirements that are not written in the authority source.
* Preserve exact numbers, ages, grades, scores, test names, AND/OR logic, and mandatory versus recommended wording.
* Minor stylistic rephrasing that does NOT change meaning is acceptable and should NOT cause a fail. Fail only for missing requirements, incorrect requirements, or meaning-changing wording. Grammar/spelling issues alone do not fail a course, but list them.

Reply with EXACTLY this JSON:
{{
  "result": "Pass" or "Fail",
  "matches_specification": true/false,
  "matches_excel": true/false,
  "missing_requirements": ["..."],
  "incorrect_requirements": ["..."],
  "wording_differences": ["..."],
  "grammar_spelling": ["..."],
  "summary": "1-3 sentence overall verdict"
}}

COURSE PAGE ENTRY REQUIREMENTS:
{page}

QUALIFICATION SPECIFICATION ENTRY REQUIREMENTS:
{spec}

EXCEL TRACKER ENTRY REQUIREMENTS:
{excel}
"""


def ai_compare(name: str, page: str, spec: str, excel: str) -> dict:
    prompt = AI_COMPARE_PROMPT.format(
        name=name,
        page=page.strip() or "(not found on the course page)",
        spec=spec.strip() or "(not available)",
        excel=excel.strip() or "(not provided)",
    )
    return parse_json_reply(call_ai(prompt, AI_COMPARE_SYSTEM))


# ═══════════════════════════════════════════════════════════════════
#  METHOD OF ASSESSMENT — compare & suggested corrected version
# ═══════════════════════════════════════════════════════════════════

AI_MOA_COMPARE_SYSTEM = (
    "You are a meticulous quality auditor for a UK college. You compare a course "
    "page's Method of Assessment against the awarding body's Qualification "
    "Specification. Reply ONLY with valid JSON — no markdown fences, no commentary."
)

AI_MOA_COMPARE_PROMPT = """Course: "{name}"

Compare the Method of Assessment from the two sources below. The QUALIFICATION SPECIFICATION is the authoritative source of truth; the COURSE PAGE is what students see and must faithfully reflect the specification.

Identify:
1. Missing assessment methods — present in the specification but absent from the course page.
2. Incorrect wording — statements on the course page that contradict or misstate the specification (e.g. wrong assessment type, wrong grading, "exam" when the spec says "assignments").
3. Additional information — content on the course page that is NOT found in the specification.
4. Grammar and formatting issues on the course page (spelling, punctuation, capitalisation, broken formatting).
5. Whether the course page wording is "identical", "similar", or "different" compared with the specification.
6. A clear overall verdict: what should be changed on the course page.

Minor stylistic rephrasing that does NOT change meaning is acceptable and should NOT cause a fail. Fail only for missing assessment methods, incorrect/contradicting statements, or additional information that misrepresents the specification.

Reply with EXACTLY this JSON:
{{
  "result": "Pass" or "Fail",
  "wording_match": "identical" or "similar" or "different",
  "missing_methods": ["..."],
  "incorrect_wording": ["..."],
  "additional_information": ["..."],
  "grammar_formatting": ["..."],
  "summary": "1-3 sentence overall verdict describing what should be changed"
}}

COURSE PAGE METHOD OF ASSESSMENT:
{page}

QUALIFICATION SPECIFICATION METHOD OF ASSESSMENT:
{spec}
"""


def ai_compare_moa(name: str, page: str, spec: str) -> dict:
    prompt = AI_MOA_COMPARE_PROMPT.format(
        name=name,
        page=page.strip() or "(not found on the course page)",
        spec=spec.strip() or "(not available)",
    )
    return parse_json_reply(call_ai(prompt, AI_MOA_COMPARE_SYSTEM))


AI_MOA_CORRECT_SYSTEM = (
    "You produce publish-ready course page copy for a UK college. "
    "Reply ONLY with the corrected text — no JSON, no code fences, no commentary "
    "before or after."
)

AI_MOA_CORRECT_PROMPT = """Suggested Corrected Method of Assessment

Extract the complete Method of Assessment section from the Qualification Specification and produce a corrected version for the course page.

Requirements:
* Use only information from the qualification specification.
* Preserve the intended meaning of the specification.
* Correct grammar, spelling, punctuation, and formatting.
* Improve readability while keeping the information accurate.
* If the course page wording differs from the specification, suggest wording that aligns with the specification.
* If information is missing from the course page, include it.
* If unnecessary information has been added to the course page, remove it.
* Present the output as clean, publish-ready text that can be copied directly into the course page.

Course: "{name}"

QUALIFICATION SPECIFICATION — METHOD OF ASSESSMENT:
{spec}

CURRENT COURSE PAGE — METHOD OF ASSESSMENT:
{page}
"""


def build_corrected_moa(qual_name: str, spec_moa: str, page_moa: str) -> str:
    """Publish-ready corrected Method of Assessment based on the spec.
    Falls back to the specification text itself if the AI reply fails or
    looks incomplete (guarding against dropped content)."""
    if not (spec_moa or "").strip():
        return ""
    fallback = spec_moa.strip()
    try:
        out = call_ai(AI_MOA_CORRECT_PROMPT.format(
            name=qual_name, spec=spec_moa.strip(),
            page=(page_moa or "").strip() or "(not found on the course page)"),
            AI_MOA_CORRECT_SYSTEM).strip()
        out = re.sub(r"^```[a-z]*\s*|\s*```$", "", out, flags=re.S).strip()
        out = re.sub(r"^suggested corrected method of assessment:?\s*", "",
                     out, flags=re.I).strip()
        if out and len(out) >= 0.5 * len(fallback):
            return out
    except Exception:
        pass
    return fallback


# ═══════════════════════════════════════════════════════════════════
#  AI-SUGGESTED ENTRY REQUIREMENTS — RELEVANT APPLICANT CRITERIA ONLY
# ═══════════════════════════════════════════════════════════════════
# Specification sections often contain surrounding policy, centre guidance,
# progression, assessment, or learner-profile prose. The website suggestion
# must not copy that whole section. The model first identifies the atomic,
# applicant-facing admission criteria and then writes concise course-page copy.
# If generation fails, the fallback is a conservative relevant-item filter —
# never the complete extracted specification section.

_BULLET_MARK = re.compile(
    r"^\s*(?:[-–—•▪●○◦*]|\d{1,2}[.)]|[a-zA-Z][.)]|o(?=\s))\s*")

_HARD_ENTRY_SIGNALS = re.compile(
    r"\b(?:minimum\s+age|aged?\s+\d|years?\s+old|entry\s+requirements?|"
    r"admission|applicants?\s+(?:must|need|require)|learners?\s+must|"
    r"must\s+(?:have|hold|be|achieve|demonstrate|provide|complete)|"
    r"required\s+to|prerequisite|prior\s+qualification|equivalent\s+qualification|"
    r"qualification|level\s*\d|grades?|GCSEs?|A[ -]?levels?|UCAS|credits?|"
    r"IELTS|CEFR|TOEFL|PTE|English\s+language|literacy|numeracy|maths?|"
    r"work\s+experience|professional\s+experience|portfolio|interview|audition|"
    r"references?|DBS|criminal\s+record|right\s+to\s+study|visa|medical|"
    r"health\s+check|alternative\s+entry|recognition\s+of\s+prior\s+learning)\b",
    re.I,
)

_UNRELATED_ENTRY_SIGNALS = re.compile(
    r"\b(?:course\s+content|curriculum|units?|learning\s+outcomes?|assessment|"
    r"grading|guided\s+learning|qualification\s+time|duration|progression|"
    r"career|employment\s+opportunities|centre(?:s)?\s+(?:must|should|are)|"
    r"provider(?:s)?\s+(?:must|should|are)|quality\s+assurance|internal\s+verification|"
    r"external\s+verification|registration|certification|resources?|delivery|"
    r"teaching|support\s+for\s+learners|reasonable\s+adjustments?|"
    r"special\s+considerations?|malpractice|appeals?|complaints?|funding|fees?)\b",
    re.I,
)

_NON_REQUIREMENT_PROFILE_SIGNALS = re.compile(
    r"\b(?:designed\s+for|intended\s+for|suitable\s+for|target\s+learners?|"
    r"recommended|desirable|ideally|would\s+benefit|should\s+ideally|"
    r"may\s+find\s+it\s+helpful)\b",
    re.I,
)


def _logical_spec_items(spec: str) -> list[str]:
    """Turn PDF/DOCX line wrapping into logical requirement-sized items."""
    text = (spec or "").strip()
    if not text:
        return []
    logical: list[str] = []
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            continue
        starts_new = bool(_BULLET_MARK.match(line))
        if logical and not starts_new and not re.search(r"[.:;!?]$", logical[-1]):
            logical[-1] += " " + line
        else:
            logical.append(line)
    items = [_BULLET_MARK.sub("", item).strip() for item in logical]
    if len(items) == 1 and len(items[0]) > 220:
        items = [part.strip() for part in
                 re.split(r"(?<=[.;])\s+(?=[A-Z(])", items[0]) if part.strip()]
    return [item for item in items if item]


def _looks_like_relevant_requirement(item: str) -> bool:
    """Conservative fallback classification used only when the AI fails."""
    item = re.sub(r"\s+", " ", item or "").strip()
    if not item or not _HARD_ENTRY_SIGNALS.search(item):
        return False
    explicit_condition = bool(re.search(
        r"\b(?:must|required|minimum|need(?:s|ed)?\s+to|shall|"
        r"eligible|admitted|accepted|entry\s+condition)\b", item, re.I))
    if _NON_REQUIREMENT_PROFILE_SIGNALS.search(item) and not explicit_condition:
        return False
    if (re.search(r"recognition\s+of\s+prior\s+learning", item, re.I)
            and not re.search(r"\b(?:entry|admission|accepted|considered|route)\b",
                              item, re.I)):
        return False
    unrelated = bool(_UNRELATED_ENTRY_SIGNALS.search(item))
    objective_value = bool(re.search(
        r"\b(?:\d+(?:\.\d+)?%?\+?|level\s*\d|grade\s*[A-Z0-9]|"
        r"IELTS|CEFR|TOEFL|PTE|GCSE|UCAS|DBS)\b", item, re.I))
    # Keep a mixed sentence only when it contains an objective applicant value;
    # otherwise surrounding administrative/policy prose is excluded.
    return not unrelated or objective_value


def relevant_requirement_items(spec: str) -> list[str]:
    """Relevant-only deterministic fallback; never returns the whole section."""
    seen, relevant = set(), []
    for item in _logical_spec_items(spec):
        if _looks_like_relevant_requirement(item):
            clean = re.sub(r"\s+", " ", item).strip(" -•\t")
            key = clean.casefold()
            if clean and key not in seen:
                seen.add(key)
                relevant.append(clean)
    return relevant


def spec_to_points(spec: str) -> str:
    """Format only conservatively relevant applicant criteria as bullets."""
    return "\n".join(f"- {item}" for item in relevant_requirement_items(spec))


def strip_excluded_sections(text: str) -> str:
    """Remove clearly unrelated lines from generated/fallback wording."""
    kept = []
    for line in (text or "").splitlines():
        clean = line.strip()
        if not clean:
            if kept and kept[-1] != "":
                kept.append("")
            continue
        # Do not remove a genuine numeric/test requirement merely because the
        # same sentence mentions an administrative concept.
        if (_UNRELATED_ENTRY_SIGNALS.search(clean)
                and not _looks_like_relevant_requirement(clean)):
            continue
        kept.append(line.rstrip())
    return re.sub(r"\n{3,}", "\n\n", "\n".join(kept)).strip()


def build_corrected_entry(spec_entry: str, ai_corrected: str) -> str:
    """Prefer relevant AI wording; fallback only to filtered applicant criteria."""
    cleaned = strip_excluded_sections((ai_corrected or "").strip())
    if cleaned:
        return cleaned
    return spec_to_points(spec_entry)


AI_FORMAT_SYSTEM = (
    "You are a meticulous UK education admissions editor. Analyse noisy "
    "qualification-document text and return only explicit applicant-facing "
    "entry criteria. Reply ONLY with valid JSON — no markdown fences or commentary."
)

AI_FORMAT_PROMPT = """Course: "{name}"

The extracted text below may include genuine Entry Requirements mixed with
neighbouring policy, administration, assessment, progression, centre guidance,
or general learner-profile information.

Your task has two stages:
1. Identify ONLY the explicit conditions an applicant must satisfy, or an
   explicitly accepted alternative route into the course.
2. Rewrite those conditions as concise, natural, publish-ready website wording.

INCLUDE ONLY WHEN EXPLICITLY STATED AS AN ADMISSION/ELIGIBILITY CONDITION:
- minimum age;
- prerequisite qualifications, levels, grades, credits, or equivalents;
- English-language, literacy, numeracy, or maths standards and exact scores;
- required experience;
- required interview, portfolio, audition, references, DBS/background check,
  visa/right-to-study, or medical/fitness check;
- an alternative entry route that the document explicitly accepts.

EXCLUDE:
- course content, units, learning outcomes, assessment, grading, duration,
  progression, careers, and marketing/background information;
- broad descriptions of the target learner or who the qualification is for,
  unless they state a testable admission condition;
- centre/provider responsibilities, registration, delivery, resources,
  quality assurance, verification, certification, or staff guidance;
- reasonable adjustments, special considerations, malpractice, appeals,
  funding, and general policy text;
- recommendations or desirable qualities that are not actual entry conditions.

WRITING RULES:
- Be concise. Do not copy the whole extracted section or add a generic summary.
- Preserve exact ages, levels, grades, scores, test names, acronyms, and AND/OR logic.
- Do not turn a recommendation into a mandatory rule, or vice versa.
- Do not invent requirements or use information from outside the supplied text.
- Use a short opening sentence only when useful, followed by clear bullets.
- Output only applicant-facing wording suitable for direct publication.

Reply with EXACTLY this JSON:
{{
  "requirements": ["one atomic, relevant entry condition per item"],
  "wording": "concise publish-ready Entry Requirements wording"
}}

EXTRACTED SPECIFICATION TEXT:
{spec}
"""


def _factual_markers(text: str) -> set[str]:
    """Objective values the concise rewrite must preserve."""
    text = text or ""
    markers = set(re.findall(
        r"(?<!\w)\d+(?:\.\d+)?(?:\s*[-–]\s*\d+(?:\.\d+)?)?%?\+?(?!\w)",
        text))
    markers.update(re.findall(
        r"(?<!\w)[A-Z]{2,}[A-Z0-9./-]*(?=s?\b)", text))
    markers.update(re.findall(r"(?<!\w)[A-Z]\d(?:[+-])?(?!\w)", text))
    return {re.sub(r"\s+", "", marker).casefold() for marker in markers}


def _ai_wording_is_complete(requirements: list[str], wording: str) -> bool:
    """Validate against the AI-selected relevant criteria, not the raw section."""
    req_text = "\n".join(requirements)
    output = strip_excluded_sections(wording)
    if not requirements or len(output) < 35:
        return False
    return _factual_markers(req_text).issubset(_factual_markers(output))


def _normalise_ai_requirements(values) -> list[str]:
    if not isinstance(values, list):
        return []
    result, seen = [], set()
    for value in values:
        item = re.sub(r"\s+", " ", str(value or "")).strip(" -•\t")
        if not item:
            continue
        # The model has already semantically selected the criteria. Apply a
        # final guard only against unmistakably unrelated output.
        if (_UNRELATED_ENTRY_SIGNALS.search(item)
                and not _looks_like_relevant_requirement(item)):
            continue
        key = item.casefold()
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


def format_spec_entry(qual_name: str, spec_entry: str) -> str:
    """Generate concise, relevant applicant-facing Entry Requirements only.

    The raw extracted section is never used as the displayed fallback. If the
    model's prose is incomplete, its filtered atomic requirement list is shown;
    if the model call fails, a conservative deterministic filter is used.
    """
    fallback = spec_to_points(spec_entry)
    if not (spec_entry or "").strip():
        return ""
    try:
        raw = call_ai(AI_FORMAT_PROMPT.format(
            name=qual_name, spec=spec_entry.strip()),
            AI_FORMAT_SYSTEM, temperature=0.2)
        payload = parse_json_reply(raw)
        requirements = _normalise_ai_requirements(payload.get("requirements"))
        wording = strip_excluded_sections(str(payload.get("wording") or "").strip())
        if _ai_wording_is_complete(requirements, wording):
            return wording
        if requirements:
            return "\n".join(f"- {item}" for item in requirements)
    except Exception:
        pass
    return fallback


def get_or_build_formatted(spec_url: str, qual_name: str, spec_entry: str) -> str:
    """Cache relevant-only wording and invalidate older full-section outputs."""
    row = get_spec(spec_url) if spec_url else None
    if (row and row.get("entry_req_formatted")
            and row.get("entry_req_format_version") == ENTRY_WORDING_VERSION):
        return strip_excluded_sections(row["entry_req_formatted"])
    formatted = format_spec_entry(qual_name, spec_entry)
    if row:
        save_spec(spec_url, entry_req_formatted=formatted,
                  entry_req_format_version=ENTRY_WORDING_VERSION)
    return formatted


# ═══════════════════════════════════════════════════════════════════
#  SPEC PROCESSING (once per unique document)
# ═══════════════════════════════════════════════════════════════════

def process_spec(url: str, force: bool = False, use_ai_fallback: bool = True) -> dict:
    """Extract & cache the Entry Requirements and Method of Assessment of one
    spec document. Skips work when the document is unchanged (hash match)
    unless force=True."""
    existing = get_spec(url)
    try:
        data = fetch_spec_bytes(url)
        doc_hash = hashlib.sha256(data).hexdigest()
        if (existing and existing["status"] == "ok" and not force
                and existing["doc_hash"] == doc_hash and existing["entry_req"]
                and existing.get("moa")
                and existing.get("extractor_version") == EXTRACTION_VERSION):
            return {"skipped": True, "status": "ok"}
        text = spec_bytes_to_text(data, url)

        entry = heuristic_entry_section(text)
        if use_ai_fallback and openrouter_api_key():
            if not entry or len(entry) < 40:
                # heuristic found nothing — let the AI search the document
                entry = ai_extract_entry(text)
            else:
                # heuristic found a section — cross-check it against an AI
                # extraction of the surrounding window; keep the more
                # complete of the two (guards against silent truncation)
                pos = text.lower().find(entry[:60].lower())
                window = text[max(0, pos - 1000):pos + len(entry) + 6000] if pos != -1 else text
                ai_entry = ai_extract_entry(window)
                if len(ai_entry) > len(entry):
                    entry = ai_entry
        if not entry:
            raise RuntimeError("Entry Requirements section not found in the document.")

        # Method of Assessment: same heuristic-plus-AI approach; a missing
        # MoA section does not fail the whole spec (ER remains the primary
        # check) — the MoA check will simply report it as unavailable.
        moa = heuristic_moa_section(text)
        if use_ai_fallback and openrouter_api_key():
            if not moa or len(moa) < 40:
                try:
                    moa = ai_extract_moa(text)
                except Exception:
                    pass
            else:
                pos = text.lower().find(moa[:60].lower())
                window = text[max(0, pos - 1000):pos + len(moa) + 6000] if pos != -1 else text
                try:
                    ai_moa = ai_extract_moa(window)
                    if len(ai_moa) > len(moa):
                        moa = ai_moa
                except Exception:
                    pass

        save_spec(url, doc_hash=doc_hash, entry_req=entry, moa=moa or "",
                  status="ok", error="", extracted_at=now(),
                  extractor_version=EXTRACTION_VERSION,
                  entry_req_formatted="", entry_req_format_version="")
        return {"skipped": False, "status": "ok"}
    except Exception as e:
        save_spec(url, status="error", error=str(e)[:500], extracted_at=now())
        return {"skipped": False, "status": "error", "error": str(e)}


# ═══════════════════════════════════════════════════════════════════
#  PDF REPORT
# ═══════════════════════════════════════════════════════════════════

@st.cache_data(show_spinner=False, max_entries=128)
def build_pdf(course, report) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle)

    issues = json.loads(report["issues_json"] or "{}")
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=16*mm, rightMargin=16*mm,
                            topMargin=16*mm, bottomMargin=16*mm,
                            title="Entry Requirements Validation Report")
    ss = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=16, spaceAfter=6)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontSize=12,
                        textColor=colors.HexColor("#333333"), spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=9.5, leading=13)

    def esc(t):
        return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") \
                        .replace("\n", "<br/>")

    story = [Paragraph("Entry Requirements Validation Report", h1)]

    result = report["result"] or "—"
    result_color = GREEN if result == "Pass" else RED
    meta = Table([
        ["Course Name", course["name"]],
        ["Course URL", course["course_url"]],
        ["Qualification Specification", course["spec_url"] or "—"],
        ["Checked", report["created_at"]],
        ["Validation Result", result],
    ], colWidths=[45*mm, 130*mm])
    meta.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F2F2")),
        ("TEXTCOLOR", (1, 4), (1, 4), colors.HexColor(result_color)),
        ("FONTNAME", (1, 4), (1, 4), "Helvetica-Bold"),
    ]))
    story += [meta, Spacer(1, 4)]

    if issues.get("summary"):
        story += [Paragraph("Summary", h2), Paragraph(esc(issues["summary"]), body)]

    story += [Paragraph("Entry Requirements — Course Page", h2),
              Paragraph(esc(report["page_entry"]) or "—", body),
              Paragraph("Entry Requirements — Qualification Specification", h2),
              Paragraph(esc(report["spec_entry"]) or "—", body),
              Paragraph("Entry Requirements — Excel Tracker", h2),
              Paragraph(esc(report["excel_entry"]) or "—", body)]

    def issue_block(title, key):
        items = issues.get(key) or []
        story.append(Paragraph(title, h2))
        if items:
            for i, it in enumerate(items, 1):
                story.append(Paragraph(f"{i}. {esc(it)}", body))
        else:
            story.append(Paragraph("None", body))

    issue_block("Missing Requirements", "missing_requirements")
    issue_block("Incorrect Requirements", "incorrect_requirements")
    issue_block("Wording Differences (meaning-changing)", "wording_differences")
    issue_block("Grammar & Spelling Issues", "grammar_spelling")

    saved_entry_wording = (report.get("corrected") or ""
                           if report.get("entry_wording_version") == ENTRY_WORDING_VERSION
                           else "")
    corrected = build_corrected_entry(report["spec_entry"], saved_entry_wording)
    story += [Paragraph("AI-Suggested Entry Requirements Wording", h2),
              Paragraph(esc(corrected) or "—", body)]

    # ── Method of Assessment check ───────────────────────────────────
    moa_issues = json.loads(report.get("moa_issues_json") or "{}")
    if report.get("spec_moa") or report.get("page_moa"):
        story += [Spacer(1, 8),
                  Paragraph("Method of Assessment Check", h1)]
        if report.get("moa_result"):
            moa_color = GREEN if report["moa_result"] == "Pass" else RED
            mt = Table([["Result", report["moa_result"]],
                        ["Wording vs specification",
                         (moa_issues.get("wording_match") or "—").capitalize()]],
                       colWidths=[45*mm, 130*mm])
            mt.setStyle(TableStyle([
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F2F2")),
                ("TEXTCOLOR", (1, 0), (1, 0), colors.HexColor(moa_color)),
                ("FONTNAME", (1, 0), (1, 0), "Helvetica-Bold"),
            ]))
            story += [mt, Spacer(1, 4)]
        if moa_issues.get("summary"):
            story += [Paragraph("Summary", h2),
                      Paragraph(esc(moa_issues["summary"]), body)]
        story += [Paragraph("Method of Assessment — Course Page", h2),
                  Paragraph(esc(report.get("page_moa")) or "—", body),
                  Paragraph("Method of Assessment — Qualification Specification", h2),
                  Paragraph(esc(report.get("spec_moa")) or "—", body)]

        def moa_block(title, key):
            items = moa_issues.get(key) or []
            story.append(Paragraph(title, h2))
            if items:
                for i, it in enumerate(items, 1):
                    story.append(Paragraph(f"{i}. {esc(it)}", body))
            else:
                story.append(Paragraph("None", body))

        moa_block("Missing Assessment Methods", "missing_methods")
        moa_block("Incorrect Wording", "incorrect_wording")
        moa_block("Additional Information (not in specification)",
                  "additional_information")
        moa_block("Grammar & Formatting Issues", "grammar_formatting")

        story += [Paragraph("Suggested Corrected Method of Assessment", h2),
                  Paragraph(esc(report.get("moa_corrected")
                                or report.get("spec_moa")) or "—", body)]

    doc.build(story)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
#  EXCEL IMPORT
# ═══════════════════════════════════════════════════════════════════

COLUMN_ALIASES = {
    "number": ["number", "no", "id", "category id"],
    "name": ["course name", "name", "course"],
    "course_url": ["course url", "url", "page url", "course link"],
    "spec_url": ["specification document", "qualification specification document url",
                 "spec url", "specification url", "qualification specification",
                 "spec document", "specification document url"],
    "level": ["level"],
    "course_type": ["type", "course type"],
    "excel_entry": ["entry requirements", "entry requirement"],
}


def import_excel(file) -> dict:
    df = pd.read_excel(file)
    cols = {str(c).strip().lower(): c for c in df.columns}

    def col(key):
        for alias in COLUMN_ALIASES[key]:
            if alias in cols:
                return cols[alias]
        return None

    name_c, url_c = col("name"), col("course_url")
    if name_c is None or url_c is None:
        raise ValueError("The Excel file must contain 'Course Name' and "
                         "'Course URL' columns.")
    mapping = {k: col(k) for k in COLUMN_ALIASES}
    imported = skipped = 0
    for _, r in df.iterrows():
        name = str(r.get(name_c, "") or "").strip()
        url = str(r.get(url_c, "") or "").strip()
        if not name or not url or not url.lower().startswith("http"):
            skipped += 1
            continue
        row = {"name": name, "course_url": url}
        for k in ("number", "spec_url", "level", "course_type", "excel_entry"):
            c = mapping.get(k)
            v = r.get(c, "") if c is not None else ""
            row[k] = "" if pd.isna(v) else str(v).strip()
        upsert_course(row)
        imported += 1
    ensure_spec_rows()
    return {"imported": imported, "skipped": skipped}


# ═══════════════════════════════════════════════════════════════════
#  CONTENT QUALITY REVIEW (original proofreading feature)
# ═══════════════════════════════════════════════════════════════════

# Keep this feature independent from the broader Course Overview Quality page.
PROOFREAD_CATEGORIES = {
    "grammar":            {"label": "Grammar",             "bg": "#FDD8D6", "border": "#E5484D"},
    "article":            {"label": "Articles (a/an/the)", "bg": "#FFE3C7", "border": "#F76B15"},
    "spelling":           {"label": "Spelling",            "bg": "#E9DDFB", "border": "#8E4EC6"},
    "punctuation":        {"label": "Punctuation & Commas","bg": "#D5E7FB", "border": "#0090FF"},
    "capitalisation":     {"label": "Capitalisation",      "bg": "#D8F3DE", "border": "#30A46C"},
    "proper_noun":        {"label": "Proper Nouns",        "bg": "#FBDCEF", "border": "#D6409F"},
    "sentence_structure": {"label": "Sentence Structure",  "bg": "#FBE8B4", "border": "#B58A00"},
    "consistency":        {"label": "Consistency",         "bg": "#D9F0F4", "border": "#0894B3"},
}

PROOFREAD_SYSTEM = (
    "You are a professional UK-English proofreader and copy editor for a college website. "
    "You review course content for grammar, articles (a/an/the), sentence structure, "
    "capitalisation, proper nouns, spelling, commas and punctuation consistency. "
    "You reply ONLY with valid JSON."
)

PROOFREAD_PROMPT = """Proofread the text below. Find every genuine error and classify it into EXACTLY one of these categories:
grammar, article, spelling, punctuation, capitalisation, proper_noun, sentence_structure, consistency

Rules:
- Return ONLY genuine errors that require a textual change.
- Do NOT return correct text, praise, observations, optional style preferences, or entries saying no change is needed.
- If a phrase is already correct, omit it completely.
- If there are no genuine errors, return an empty "issues" list.
- "original" must be an EXACT substring copied verbatim from the text (short — the smallest span that contains the problem).
- "correction" must be the fixed version of that span and must differ from "original".
- "explanation" is one short sentence explaining the actual error.
- Set "has_issue" to true for every returned issue.
- Also produce the fully corrected version of the whole text. Do not make unlisted changes.

Reply with EXACTLY this JSON:
{{
  "issues": [
    {{"has_issue": true, "category": "...", "original": "...", "correction": "...", "explanation": "..."}}
  ],
  "corrected_text": "..."
}}

=== TEXT TO REVIEW ===
{text}
"""


_NO_ISSUE_PHRASES = (
    "no change needed",
    "no changes needed",
    "no correction needed",
    "no correction required",
    "no edit needed",
    "already correct",
    "is correct",
    "correctly capitalised",
    "correctly capitalized",
    "no issue",
    "not an issue",
)


def _normalise_proofread_span(value: object) -> str:
    """Normalise invisible spacing without hiding real case or punctuation changes."""
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = text.replace("\u00a0", " ")
    return " ".join(text.split()).strip()


def _is_genuine_proofread_issue(issue: object, source_text: str) -> bool:
    """Reject AI observations and no-op edits before they reach the UI or score."""
    if not isinstance(issue, dict):
        return False

    # Newer prompt responses include this field, but remain compatible with
    # earlier model responses that omit it.
    if issue.get("has_issue") is False:
        return False

    category = str(issue.get("category", "")).strip()
    original_raw = str(issue.get("original", ""))
    correction_raw = str(issue.get("correction", ""))
    explanation = _normalise_proofread_span(issue.get("explanation", "")).lower()

    if category not in PROOFREAD_CATEGORIES:
        return False
    if not original_raw.strip() or not correction_raw.strip():
        return False
    if original_raw not in source_text:
        return False

    original = _normalise_proofread_span(original_raw)
    correction = _normalise_proofread_span(correction_raw)
    if not original or not correction or original == correction:
        return False
    if any(phrase in explanation for phrase in _NO_ISSUE_PHRASES):
        return False
    return True


def _apply_proofread_corrections(source_text: str, issues: list[dict]) -> str:
    """Build corrected copy using only validated issues, preventing silent edits."""
    corrected = source_text
    for issue in issues:
        original = str(issue.get("original", ""))
        replacement = str(issue.get("correction", ""))
        if original and original in corrected:
            corrected = corrected.replace(original, replacement, 1)
    return corrected


def normalise_proofreading_result(raw: object, source_text: str) -> dict:
    """Sanitise the model response and remove duplicate or false-positive issues."""
    raw = raw if isinstance(raw, dict) else {}
    candidates = raw.get("issues", [])
    if not isinstance(candidates, list):
        candidates = []

    issues: list[dict] = []
    seen: set[tuple[str, str, str]] = set()
    for item in candidates:
        if not _is_genuine_proofread_issue(item, source_text):
            continue
        cleaned = {
            "category": str(item.get("category", "")).strip(),
            "original": str(item.get("original", "")),
            "correction": str(item.get("correction", "")),
            "explanation": str(item.get("explanation", "")).strip(),
        }
        identity = (
            cleaned["category"],
            cleaned["original"],
            cleaned["correction"],
        )
        if identity in seen:
            continue
        seen.add(identity)
        issues.append(cleaned)

    return {
        "issues": issues,
        "corrected_text": _apply_proofread_corrections(source_text, issues),
    }


def run_proofreading_review(text: str) -> dict:
    raw = parse_json_reply(call_ai(PROOFREAD_PROMPT.format(text=text), PROOFREAD_SYSTEM))
    return normalise_proofreading_result(raw, text)


def annotate_proofreading_html(text: str, issues: list) -> str:
    """Return proofreading markup using the original eight categories."""
    escaped = html.escape(text)
    for n, issue in enumerate(issues, start=1):
        original = html.escape(str(issue.get("original", "")))
        if not original:
            continue
        category = issue.get("category", "grammar")
        style = PROOFREAD_CATEGORIES.get(category, PROOFREAD_CATEGORIES["grammar"])
        tip = html.escape(
            f"{style['label']}: {issue.get('correction', '')} — "
            f"{issue.get('explanation', '')}"
        )
        mark = (
            f'<mark class="qr-mark" style="background:{style["bg"]};'
            f'border-bottom:2px solid {style["border"]};" title="{tip}">'
            f'<sup class="qr-num" style="background:{style["border"]};">{n}</sup>{original}</mark>'
        )
        escaped = escaped.replace(original, mark, 1)
    return escaped.replace("\n", "<br>")


# ═══════════════════════════════════════════════════════════════════
#  COURSE OVERVIEW QUALITY REVIEW
# ═══════════════════════════════════════════════════════════════════

# The AI scores each editorial criterion from 0-100. The overall score is the
# simple average of the criteria; no source article or external standard assigns
# percentage weights to these checks.
QUALITY_CRITERIA = {
    "professional_tone": {
        "label": "Professional tone", "bg": "#E9DDFB", "border": "#8E4EC6",
    },
    "appropriate_formality": {
        "label": "Appropriate formality", "bg": "#FBDCEF", "border": "#D6409F",
    },
    "clarity": {
        "label": "Clarity", "bg": "#D5E7FB", "border": "#0090FF",
    },
    "readability": {
        "label": "Readability & plain language", "bg": "#D9F0F4", "border": "#0894B3",
    },
    "learner_focus": {
        "label": "Learner-focused writing", "bg": "#FFE3C7", "border": "#F76B15",
    },
    "course_purpose": {
        "label": "Clear course purpose", "bg": "#FBE8B4", "border": "#B58A00",
    },
    "learner_benefits": {
        "label": "Clear learner benefits", "bg": "#D7F0E5", "border": "#218358",
    },
    "structure_flow": {
        "label": "Structure & flow", "bg": "#D8F3DE", "border": "#30A46C",
    },
    "conciseness_repetition": {
        "label": "Conciseness & repetition", "bg": "#E8E8E8", "border": "#6F6F6F",
    },
    "language_accuracy": {
        "label": "Grammar & sentence quality", "bg": "#FDD8D6", "border": "#E5484D",
    },
    "consistency_uk_english": {
        "label": "Consistency & UK English", "bg": "#E1E7FF", "border": "#5B5BD6",
    },
    "promotional_balance": {
        "label": "Promotional balance & specificity", "bg": "#F9E1C7", "border": "#C25D05",
    },
}


QUALITY_SYSTEM = (
    "You are a senior UK web-content editor. Review only the writing quality "
    "of a college course overview. Do not fact-check qualification claims, "
    "accreditation, credits, duration, progression, careers or other course facts. "
    "Reply ONLY with valid JSON."
)

QUALITY_PROMPT = """Evaluate only the writing quality of the COURSE OVERVIEW below against all twelve editorial criteria.

EDITORIAL CRITERIA
- professional_tone: professional, confident, approachable and suitable for a college website
- appropriate_formality: neither too casual nor unnecessarily academic, stiff or elaborate
- clarity: precise, specific and easy to understand without ambiguous or vague wording
- readability: plain UK English, manageable sentence length, limited jargon and smooth reading
- learner_focus: speaks to the prospective learner and addresses their interests rather than focusing mainly on the provider
- course_purpose: quickly and clearly explains what the course is and what it covers
- learner_benefits: clearly communicates useful knowledge, skills, outcomes or progression value for the learner
- structure_flow: logical paragraph order, effective opening, coherent transitions and a suitable ending
- conciseness_repetition: avoids repeated ideas, repeated words, filler and unnecessary wording
- language_accuracy: sound grammar, spelling, punctuation, articles and sentence construction
- consistency_uk_english: consistent terminology, capitalisation, hyphenation, point of view and UK English
- promotional_balance: persuasive but not exaggerated, generic, sales-heavy or filled with unsupported-sounding superlatives

IMPORTANT SCOPE RULES
1. Review only how the Course Overview is written.
2. Do NOT check whether factual statements are true or ask for evidence, sources or human review.
3. Do NOT assess entry requirements, assessment methods or other webpage sections.
4. Score EACH criterion from 0 to 100. The application calculates the overall score as an equal average.
5. A high score requires positive evidence in the wording, not merely the absence of errors.
6. For every issue, "original" MUST be an exact, short substring copied verbatim from the overview.
7. Use only these severity values: major, minor, suggestion.
8. Use only the exact criterion keys listed above.
9. The corrected_text must preserve all supplied course facts and meaning. Improve only the wording, organisation and readability.
10. Formality is not automatically good: prefer professional, natural and approachable wording.
11. Do not add new course facts, accreditation, salary, duration, entry, assessment or career claims.
12. Return no markdown and no text outside the JSON.

COURSE TITLE CONTEXT
{context}

AUTOMATIC TEXT METRICS
{metrics}

Return EXACTLY this JSON shape:
{{
  "summary": "Two or three sentences explaining the overall writing quality.",
  "category_scores": {{
    "professional_tone": {{"score": 0, "explanation": "..."}},
    "appropriate_formality": {{"score": 0, "explanation": "..."}},
    "clarity": {{"score": 0, "explanation": "..."}},
    "readability": {{"score": 0, "explanation": "..."}},
    "learner_focus": {{"score": 0, "explanation": "..."}},
    "course_purpose": {{"score": 0, "explanation": "..."}},
    "learner_benefits": {{"score": 0, "explanation": "..."}},
    "structure_flow": {{"score": 0, "explanation": "..."}},
    "conciseness_repetition": {{"score": 0, "explanation": "..."}},
    "language_accuracy": {{"score": 0, "explanation": "..."}},
    "consistency_uk_english": {{"score": 0, "explanation": "..."}},
    "promotional_balance": {{"score": 0, "explanation": "..."}}
  }},
  "strengths": ["..."],
  "issues": [
    {{
      "criterion": "clarity",
      "severity": "major",
      "original": "exact substring",
      "correction": "improved replacement",
      "explanation": "Why this wording reduces the quality of the overview."
    }}
  ],
  "corrected_text": "Complete improved version of the supplied Course Overview."
}}

=== COURSE OVERVIEW TO EVALUATE ===
{text}
"""


def content_metrics(text: str) -> dict:
    """Return small, deterministic readability indicators for the UI and AI."""
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    words = re.findall(r"\b[\w'-]+\b", cleaned)
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", cleaned) if s.strip()]
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text or "") if p.strip()]
    sentence_lengths = [len(re.findall(r"\b[\w'-]+\b", s)) for s in sentences]
    paragraph_lengths = [len(re.findall(r"\b[\w'-]+\b", p)) for p in paragraphs]
    return {
        "word_count": len(words),
        "sentence_count": len(sentences),
        "paragraph_count": len(paragraphs),
        "average_sentence_words": round(sum(sentence_lengths) / max(1, len(sentence_lengths)), 1),
        "long_sentences_25_plus": sum(1 for n in sentence_lengths if n >= 25),
        "average_paragraph_words": round(sum(paragraph_lengths) / max(1, len(paragraph_lengths)), 1),
        "estimated_reading_minutes": max(1, round(len(words) / 220)) if words else 0,
    }


def _safe_score(value, default=0) -> int:
    try:
        return max(0, min(100, int(round(float(value)))))
    except (TypeError, ValueError):
        return default


def normalise_quality_result(raw: dict, original_text: str) -> dict:
    """Normalise model JSON and calculate an equal-average editorial score."""
    raw = raw if isinstance(raw, dict) else {}
    supplied = raw.get("category_scores") or {}
    if not isinstance(supplied, dict):
        supplied = {}
    categories = {}
    score_values = []
    for key in QUALITY_CRITERIA:
        item = supplied.get(key) or {}
        if isinstance(item, (int, float, str)):
            item = {"score": item, "explanation": ""}
        score = _safe_score(item.get("score"), 0)
        categories[key] = {
            "score": score,
            "explanation": str(item.get("explanation") or "").strip(),
        }
        score_values.append(score)

    issues = []
    valid_severity = {"major", "minor", "suggestion"}
    for issue in raw.get("issues") or []:
        if not isinstance(issue, dict):
            continue
        criterion = str(issue.get("criterion") or "language_accuracy").strip()
        if criterion not in QUALITY_CRITERIA:
            criterion = "language_accuracy"
        severity = str(issue.get("severity") or "minor").strip().lower()
        if severity not in valid_severity:
            severity = "minor"
        original = str(issue.get("original") or "").strip()
        if original and original not in original_text:
            original = ""
        issues.append({
            "criterion": criterion,
            "severity": severity,
            "original": original,
            "correction": str(issue.get("correction") or "").strip(),
            "explanation": str(issue.get("explanation") or "").strip(),
        })

    score = int(round(sum(score_values) / max(1, len(score_values))))
    if score >= 80:
        status = "Pass"
    elif score >= 65:
        status = "Needs Improvement"
    else:
        status = "Fail"

    raw_strengths = raw.get("strengths") or []
    if isinstance(raw_strengths, str):
        raw_strengths = [raw_strengths]
    return {
        "overall_score": score,
        "status": status,
        "summary": str(raw.get("summary") or "").strip(),
        "category_scores": categories,
        "strengths": [str(x).strip() for x in raw_strengths if str(x).strip()],
        "issues": issues,
        "corrected_text": str(raw.get("corrected_text") or original_text).strip(),
    }


def run_quality_review(text: str, context: str = "No additional course context supplied.") -> dict:
    metrics = content_metrics(text)
    raw = parse_json_reply(call_ai(
        QUALITY_PROMPT.format(
            text=text[:18000],
            context=context[:5000],
            metrics=json.dumps(metrics, ensure_ascii=False),
        ),
        QUALITY_SYSTEM,
    ))
    return normalise_quality_result(raw, text)


OVERVIEW_EXTRACT_PROMPT = """Extract only the main course overview or course-description copy from the webpage text below.
Exclude navigation, pricing, reviews, FAQs, entry requirements, assessment, course curriculum, careers, contact details and footer text.
Preserve the overview wording exactly as it appears. Do not rewrite or summarise it.
Reply only with valid JSON in this shape: {{"overview": "..."}}

WEBPAGE TEXT:
{text}
"""


def ai_extract_overview(page_text: str) -> str:
    raw = parse_json_reply(call_ai(
        OVERVIEW_EXTRACT_PROMPT.format(text=page_text[:14000]),
        "You extract an exact course-overview section from webpage text and reply only with JSON.",
    ))
    return str(raw.get("overview") or "").strip()[:6000]


def extract_course_overview_section(text: str) -> str:
    """Return only the Course Overview section from pasted or uploaded text.

    When no overview heading is present, the supplied text is treated as the
    overview itself. This supports users who paste only the section body.
    """
    source = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not source:
        return ""
    stop_headings = _stop_re(PAGE_NEXT_AFTER_OVERVIEW)
    overview = heuristic_section(
        source,
        r"course\s+overview",
        stop_headings,
        max_chars=12000,
    )
    if not overview:
        overview = heuristic_section(
            source,
            OVERVIEW_HEADING,
            stop_headings,
            max_chars=12000,
        )
    return (overview or source).strip()[:12000]


def extract_uploaded_content(file) -> str:
    """Extract text from an uploaded TXT, DOCX or text-based PDF."""
    if file is None:
        return ""
    data = file.getvalue()
    name = (getattr(file, "name", "") or "").lower()
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="ignore").strip()
    if name.endswith(".docx"):
        from docx import Document
        document = Document(io.BytesIO(data))
        parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                line = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if line:
                    parts.append(line)
        return "\n".join(parts).strip()
    if name.endswith(".pdf"):
        text_parts = []
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text.strip())
        except Exception:
            text_parts = []
        if not text_parts:
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(data))
                text_parts = [(page.extract_text() or "").strip() for page in reader.pages]
            except Exception as exc:
                raise RuntimeError(f"The PDF could not be read: {exc}") from exc
        text = "\n\n".join(x for x in text_parts if x).strip()
        if not text:
            raise RuntimeError(
                "No selectable text was found in the PDF. Upload a text-based PDF or DOCX."
            )
        return text
    raise RuntimeError("Unsupported document type. Upload TXT, MD, DOCX or PDF.")


def course_quality_context(course: dict) -> str:
    if not course:
        return "No course title was supplied."
    lines = [f"Course title: {course.get('name') or 'Not supplied'}"]
    if course.get("level"):
        lines.append(f"Level: {course.get('level')}")
    if course.get("course_type"):
        lines.append(f"Course type: {course.get('course_type')}")
    return "\n".join(lines)


def annotate_text_html(text: str, issues: list) -> str:
    """Return HTML with source issues highlighted and numbered."""
    escaped = html.escape(text)
    for n, issue in enumerate(issues, start=1):
        original = html.escape(str(issue.get("original", "")))
        if not original:
            continue
        criterion = issue.get("criterion", "language_accuracy")
        style = QUALITY_CRITERIA.get(criterion, QUALITY_CRITERIA["language_accuracy"])
        tip = html.escape(
            f"{style['label']} · {issue.get('severity', 'minor').title()}: "
            f"{issue.get('correction', '')} — {issue.get('explanation', '')}"
        )
        mark = (
            f'<mark class="qr-mark" style="background:{style["bg"]};'
            f'border-bottom:2px solid {style["border"]};" title="{tip}">'
            f'<sup class="qr-num" style="background:{style["border"]};">{n}</sup>{original}</mark>'
        )
        escaped = escaped.replace(original, mark, 1)
    return escaped.replace("\n", "<br>")


QR_CSS = """
<style>
.qr-paper {
  background:#FFFFFF; border:1px solid #D9E2EC; border-radius:16px;
  padding:26px 30px; line-height:2.05; font-size:1.0rem; color:#2A2F3A;
  box-shadow: 0 8px 24px rgba(16,42,67,.07);
}
.qr-mark { border-radius:4px; padding:1px 3px; cursor:help; position:relative; }
.qr-num {
  color:#fff; font-size:.62rem; font-weight:700; border-radius:999px;
  padding:0 4px; margin-right:2px; position:relative; top:-7px;
}
.qr-legend span {
  display:inline-block; margin:3px 8px 3px 0; padding:3px 10px; border-radius:999px;
  font-size:.76rem; font-weight:600;
}
.issue-card {
  border-radius:13px; border:1px solid #D9E2EC; border-left-width:5px;
  padding:12px 16px; margin-bottom:10px; background:#fff;
}
.issue-card .cat { font-size:.72rem; font-weight:700; text-transform:uppercase; letter-spacing:.06em;}
.issue-card .orig { text-decoration:line-through; color:#B0341F; }
.issue-card .corr { color:#1D7A46; font-weight:600; }
.quality-score-row {
  display:flex; justify-content:space-between; gap:16px; align-items:flex-start;
  border-bottom:1px solid #E7EDF3; padding:11px 0;
}
.quality-score-row:last-child { border-bottom:0; }
.quality-score-label { font-weight:700; color:#243B53; }
.quality-score-explanation { color:#627D98; font-size:.88rem; margin-top:2px; }
.quality-score-number { font-size:1.05rem; font-weight:800; white-space:nowrap; }
</style>
"""


def qr_stat(col, value, label, kind="info"):
    cards.stat_card(col, value, label, kind)


# ═══════════════════════════════════════════════════════════════════
#  UI
# ═══════════════════════════════════════════════════════════════════

st.set_page_config(page_title="Course Content Checker", page_icon="✅",
                   layout="wide", initial_sidebar_state="expanded")
init_db()
styles.inject()

# Open the application on Run Check on the first session load.
# The value is only seeded once, so subsequent sidebar navigation remains
# fully user-controlled across Streamlit reruns.
if "main_navigation" not in st.session_state:
    st.session_state["main_navigation"] = "▶️ Run Check"

page = app_sidebar.render(APP_VERSION, bool(openrouter_api_key()))


def badge(result: str) -> str:
    return cards.badge(result or "—")


# ── PAGE: Upload & Specs ────────────────────────────────────────────
if page == "📥 Upload & Specs":
    header.page_header("📥 Upload & Specs",
                       "Import the tracker Excel and extract each qualification "
                       "specification once — everything is cached locally.",
                       chip="Data preparation")
    st.subheader("Upload Tracker Excel")
    upload_ui.info(".xlsx · .xlsm · .xls",
                   "Columns used: Course Name · Course URL · Specification Document "
                   "· Level · Type · Entry Requirements (auto-parsed when missing). "
                   "Data is stored until removed in 🗂️ Manage Data.")
    up = st.file_uploader("Drag and drop your tracker here",
                          type=["xlsx", "xlsm", "xls"])
    upload_ui.file_details(up)
    if up is not None and st.button("📥 Import courses", type="primary"):
        try:
            res = import_excel(up)
            st.success(f"✅ Import complete — {res['imported']} courses "
                       f"imported/updated, {res['skipped']} rows skipped.")
        except Exception as e:
            st.error(f"Import failed: {e}")

    courses = all_courses()
    if courses:
        st.subheader(f"Courses stored ({len(courses)})")
        st.dataframe(pd.DataFrame(
            [{"Number": c["number"], "Course Name": c["name"],
              "Level": c["level"], "Type": c["course_type"],
              "Spec URL": c["spec_url"]} for c in courses]),
            width='stretch', height=280)

    st.divider()
    st.header("Qualification Specification Documents")
    ensure_spec_rows()
    specs = all_specs()
    if not specs:
        st.info("Import the Excel first — unique specification documents will be "
                "listed here.")
    else:
        pending = [s for s in specs if s["status"] != "ok"]
        st.caption(f"{len(specs)} unique documents · "
                   f"{len(specs) - len(pending)} extracted · {len(pending)} pending/error. "
                   "Each document is processed once and cached; it is only "
                   "re-extracted if the file changes or you force it.")
        c1, c2 = st.columns(2)
        run_pending = c1.button("Process pending documents", type="primary",
                                disabled=not pending)
        run_all = c2.button("Re-check all (skips unchanged files)")
        if run_pending or run_all:
            targets = specs if run_all else pending
            prog = progress_ui.BulkProgress(len(targets))
            for s in targets:
                process_spec(s["url"])
                prog.update(s["url"])
            prog.finish()
            st.rerun()

        current_specs = all_specs()
        sp1, sp2 = st.columns([1, 1])
        upload_spec_page_size = sp1.selectbox(
            "Specifications per page", [10, 25, 50],
            key="upload_spec_page_size")
        upload_spec_page_count = max(
            1, (len(current_specs) + upload_spec_page_size - 1)
            // upload_spec_page_size)
        if not 1 <= int(st.session_state.get(
                "upload_spec_page", 1)) <= upload_spec_page_count:
            st.session_state["upload_spec_page"] = 1
        upload_spec_page = int(sp2.number_input(
            "Specification page", min_value=1,
            max_value=upload_spec_page_count, value=1, step=1,
            key="upload_spec_page"))
        upload_spec_start = (upload_spec_page - 1) * upload_spec_page_size
        visible_upload_specs = current_specs[
            upload_spec_start:upload_spec_start + upload_spec_page_size]
        st.caption(
            f"Showing {upload_spec_start + 1}–"
            f"{upload_spec_start + len(visible_upload_specs)} of "
            f"{len(current_specs)} specifications.")

        for s in visible_upload_specs:
            n_courses = get_conn().execute(
                "SELECT COUNT(*) c FROM courses WHERE spec_url=?",
                (s["url"],)).fetchone()["c"]
            icon = {"ok": "🟢", "error": "🔴"}.get(s["status"], "⚪")
            with st.expander(f"{icon} {s['url']}  ·  used by {n_courses} course(s)"):
                if s["status"] == "error":
                    st.error(s["error"])
                edited = st.text_area("Entry Requirements (cached — editable)",
                                      value=s["entry_req"] or "", height=160,
                                      key=f"spec_{s['id']}")
                edited_moa = st.text_area("Method of Assessment (cached — editable)",
                                          value=s.get("moa") or "", height=160,
                                          key=f"specmoa_{s['id']}")
                b1, b2 = st.columns(2)
                if b1.button("Save edits", key=f"save_{s['id']}"):
                    save_spec(s["url"], entry_req=edited, moa=edited_moa,
                              status="ok", error="",
                              extracted_at=now(), entry_req_formatted="",
                              entry_req_format_version="")
                    st.success("Saved.")
                if b2.button("Force re-extract", key=f"re_{s['id']}"):
                    with st.spinner("Re-extracting…"):
                        process_spec(s["url"], force=True)
                    st.rerun()


# ── PAGE: Run Check ─────────────────────────────────────────────────
elif page == "▶️ Run Check":
    header.page_header("▶️ Run Check",
                       "Audit a course page against its qualification "
                       "specification — Entry Requirements and Method of "
                       "Assessment in one pass.",
                       chip="Compliance audit")
    courses = all_courses()
    if not courses:
        st.info("Upload the tracker Excel first (📥 Upload & Specs).")
        st.stop()

    levels = sorted({c["level"] for c in courses if c["level"]})
    types = sorted({c["course_type"] for c in courses if c["course_type"]})
    f1, f2, f3 = st.columns([1, 1, 2])
    f_level = f1.selectbox("Level", ["All"] + levels)
    f_type = f2.selectbox("Type", ["All"] + types)
    f_text = f3.text_input("Search course name")

    filtered = [c for c in courses
                if (f_level == "All" or c["level"] == f_level)
                and (f_type == "All" or c["course_type"] == f_type)
                and (f_text.lower() in c["name"].lower())]
    if not filtered:
        st.warning("No courses match the filters.")
        st.stop()

    sel = st.selectbox("Course", filtered,
                       format_func=lambda c: f"{c['number']} — {c['name']}"
                       if c["number"] else c["name"])
    course = get_course(sel["id"])

    st.markdown(f"**Course URL:** {course['course_url']}  \n"
                f"**Specification:** {course['spec_url'] or '—'}")

    spec_row = get_spec(course["spec_url"]) if course["spec_url"] else None
    spec_ready = bool(spec_row and spec_row["status"] == "ok"
                      and spec_row["entry_req"]
                      and spec_row.get("extractor_version") == EXTRACTION_VERSION)
    if course["spec_url"] and not spec_ready:
        st.warning("This course's specification hasn't been extracted yet — it will "
                   "be processed automatically when you run the check (and cached "
                   "for every course sharing it).")
    elif not course["spec_url"]:
        st.warning("No specification document URL for this course — the check will "
                   "compare the page against the Excel tracker only.")

    api_key_ready = bool(openrouter_api_key())
    if not api_key_ready:
        st.error(
            "Run Check is unavailable because `OPENROUTER_API_KEY` is not "
            "configured. For local use, copy `.streamlit/secrets.toml.example` "
            "to `.streamlit/secrets.toml` and replace the placeholder key. "
            "On Streamlit Cloud, add the same setting under App settings → "
            "Secrets."
        )

    run_check = st.button(
        "▶️ Run Check",
        type="primary",
        key="run_check_button",
        disabled=not api_key_ready,
        help=(None if api_key_ready else
              "Configure OPENROUTER_API_KEY before running a check."),
    )
    if run_check:
        try:
            prog = progress_ui.StepProgress(4, course["name"])
            prog.step(1, "Reading course page")
            with st.spinner("Reading course page…"):
                page_entry, page_moa, full_text = extract_page_sections(course["course_url"])
                # Cross-check the page section. This catches accordions/CMS
                # layouts where a deterministic extraction returns only the
                # first requirement or stops at a misleading short line.
                try:
                    ai_page_entry = ai_extract_entry(
                        section_ai_window(full_text, ENTRY_HEADING))
                    page_entry = reconcile_section(page_entry, ai_page_entry)
                except Exception:
                    # The main comparison still runs with the deterministic
                    # extraction if this optional cross-check fails.
                    pass
                if not page_entry:
                    raise RuntimeError(
                        "No Entry Requirements section could be extracted from "
                        "the course page. Check that the page is publicly "
                        "accessible and contains an Entry Requirements heading.")
                try:
                    ai_page_moa = ai_extract_moa(
                        section_ai_window(full_text, MOA_HEADING))
                    page_moa = reconcile_section(page_moa, ai_page_moa)
                except Exception:
                    page_moa = page_moa or ""
            spec_entry = spec_moa = ""
            spec_row = None
            cached_entry_wording = ""
            prog.step(2, "Loading specification (cached when possible)")
            if course["spec_url"]:
                with st.spinner("Loading specification…"):
                    if not spec_ready:
                        process_spec(course["spec_url"])
                    spec_row = get_spec(course["spec_url"])
                    if spec_row and spec_row["status"] == "ok":
                        spec_entry = spec_row["entry_req"] or ""
                        spec_moa = spec_row.get("moa") or ""
                        if not spec_moa:  # cached before MoA support → refresh
                            process_spec(course["spec_url"], force=True)
                            spec_row = get_spec(course["spec_url"])
                            spec_moa = (spec_row.get("moa") or "") if spec_row else ""
                            spec_entry = (spec_row.get("entry_req") or spec_entry) if spec_row else spec_entry
                    else:
                        st.warning("Specification could not be extracted: "
                                   f"{spec_row['error'] if spec_row else 'unknown error'}")
            if (spec_row and spec_row.get("entry_req_formatted")
                    and spec_row.get("entry_req_format_version") == ENTRY_WORDING_VERSION):
                cached_entry_wording = strip_excluded_sections(
                    spec_row["entry_req_formatted"])
            # ── steps 3+4: all AI checks run IN PARALLEL ──────────────
            prog.step(3, "Running AI checks (Entry Requirements + Method of "
                         "Assessment in parallel)")
            # worker threads can't always read st.secrets — mirror the key
            # into the environment so call_ai's fallback finds it
            _k = openrouter_api_key()
            if _k:
                os.environ["OPENROUTER_API_KEY"] = _k

            with concurrent.futures.ThreadPoolExecutor(max_workers=4) as pool:
                f_er = pool.submit(ai_compare, course["name"], page_entry,
                                   spec_entry, course["excel_entry"] or "")
                f_entry_wording = (pool.submit(format_spec_entry,
                                               course["name"], spec_entry)
                                   if spec_entry.strip() and not cached_entry_wording
                                   else None)
                f_moa = f_corr_moa = None
                if spec_moa.strip() or page_moa.strip():
                    f_moa = pool.submit(ai_compare_moa, course["name"],
                                        page_moa, spec_moa)
                    f_corr_moa = pool.submit(build_corrected_moa,
                                             course["name"], spec_moa, page_moa)
                verdict = f_er.result()
                ai_entry_wording = (f_entry_wording.result()
                                    if f_entry_wording else cached_entry_wording)
                moa_verdict = f_moa.result() if f_moa else {}
                moa_corrected = f_corr_moa.result() if f_corr_moa else ""

            if (f_entry_wording and spec_row and ai_entry_wording):
                save_spec(course["spec_url"],
                          entry_req_formatted=ai_entry_wording,
                          entry_req_format_version=ENTRY_WORDING_VERSION)

            prog.step(4, "Finalising report")
            corrected = build_corrected_entry(spec_entry, ai_entry_wording)
            moa_result = ""
            if moa_verdict:
                moa_result = ("Pass" if str(moa_verdict.get("result", "")).lower()
                              == "pass" else "Fail")
            result = "Pass" if str(verdict.get("result", "")).lower() == "pass" else "Fail"
            save_report(course["id"], result, page_entry, spec_entry,
                        course["excel_entry"] or "", verdict, corrected,
                        page_moa=page_moa, spec_moa=spec_moa,
                        moa_issues=moa_verdict, moa_corrected=moa_corrected,
                        moa_result=moa_result)
            prog.done("Check complete")
            st.success("✅ Check complete — report saved.")
        except Exception as e:
            st.error(f"Check failed: {e}")

    report = latest_report(course["id"])
    if report:
        issues = json.loads(report["issues_json"] or "{}")
        moa_issues = json.loads(report.get("moa_issues_json") or "{}")
        st.divider()
        st.subheader("Validation Report")
        moa_badge_html = (f" &nbsp;·&nbsp; **Method of Assessment:** "
                          f"{badge(report['moa_result'])}"
                          if report.get("moa_result") else "")
        st.markdown(f"**Entry Requirements:** {badge(report['result'])}"
                    f"{moa_badge_html} &nbsp;·&nbsp; "
                    f"checked {report['created_at']}", unsafe_allow_html=True)

        tab_er, tab_moa = st.tabs(["📋 Entry Requirements",
                                   "📝 Method of Assessment"])

        def show_issues(issue_dict, title, key, color, prefix):
            items = issue_dict.get(key) or []
            st.markdown(f"**{title}** "
                        f"<span style='color:{color}'>({len(items)})</span>",
                        unsafe_allow_html=True)
            if items:
                for it in items:
                    st.markdown(f"- {it}")
            else:
                st.caption("None")

        # ── TAB 1: Entry Requirements ────────────────────────────────
        with tab_er:
            if issues.get("summary"):
                st.markdown(f"> {issues['summary']}")

            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown("**Course Page**")
                st.text_area("page", report["page_entry"] or "—", height=220,
                             label_visibility="collapsed", disabled=True)
            with c2:
                st.markdown("**Qualification Specification**")
                st.text_area("spec", report["spec_entry"] or "—", height=220,
                             label_visibility="collapsed", disabled=True)
            with c3:
                st.markdown("**Excel Tracker**")
                st.text_area("excel", report["excel_entry"] or "—", height=220,
                             label_visibility="collapsed", disabled=True)

            i1, i2 = st.columns(2)
            with i1:
                show_issues(issues, "Missing Requirements",
                            "missing_requirements", RED, "er")
                show_issues(issues, "Incorrect Requirements",
                            "incorrect_requirements", RED, "er")
            with i2:
                show_issues(issues, "Wording Differences",
                            "wording_differences", AMBER, "er")
                show_issues(issues, "Grammar & Spelling",
                            "grammar_spelling", AMBER, "er")

            st.markdown("**✨ AI-Suggested Entry Requirements Wording**")
            st.caption("Concise, applicant-facing wording generated only from the "
                       "relevant admission criteria in the qualification specification. "
                       "Review before publishing.")
            saved_entry_wording = (report.get("corrected") or ""
                                   if report.get("entry_wording_version") == ENTRY_WORDING_VERSION
                                   else "")
            corrected_txt = build_corrected_entry(
                report["spec_entry"], saved_entry_wording)
            st.markdown(corrected_txt or "—")

        # ── TAB 2: Method of Assessment ──────────────────────────────
        with tab_moa:
            if not (report.get("spec_moa") or report.get("page_moa")):
                st.info("No Method of Assessment was found — run the check "
                        "again to extract it (older reports predate this "
                        "check).")
            else:
                if moa_issues.get("summary"):
                    st.markdown(f"> {moa_issues['summary']}")
                wm = (moa_issues.get("wording_match") or "").capitalize()
                if wm:
                    wm_color = {"Identical": GREEN, "Similar": AMBER,
                                "Different": RED}.get(wm, AMBER)
                    st.markdown(f"**Wording vs specification:** "
                                f"<span style='color:{wm_color};font-weight:700'>"
                                f"{wm}</span>", unsafe_allow_html=True)

                m1, m2 = st.columns(2)
                with m1:
                    st.markdown("**Course Page**")
                    st.text_area("moa_page", report.get("page_moa") or "—",
                                 height=220, label_visibility="collapsed",
                                 disabled=True)
                with m2:
                    st.markdown("**Qualification Specification**")
                    st.text_area("moa_spec", report.get("spec_moa") or "—",
                                 height=220, label_visibility="collapsed",
                                 disabled=True)

                j1, j2 = st.columns(2)
                with j1:
                    show_issues(moa_issues, "Missing Assessment Methods",
                                "missing_methods", RED, "moa")
                    show_issues(moa_issues, "Incorrect Wording",
                                "incorrect_wording", RED, "moa")
                with j2:
                    show_issues(moa_issues, "Additional Information "
                                "(not in specification)",
                                "additional_information", AMBER, "moa")
                    show_issues(moa_issues, "Grammar & Formatting",
                                "grammar_formatting", AMBER, "moa")

                st.markdown("**Suggested Corrected Method of Assessment**")
                st.caption("Publish-ready text based only on the qualification "
                           "specification — can be copied directly into the "
                           "course page.")
                st.markdown(report.get("moa_corrected")
                            or (report.get("spec_moa") or "—"))

        pdf = build_pdf(course, report)
        st.download_button("⬇️ Download Report (PDF)", data=pdf,
                           file_name=f"ER_Report_{re.sub(r'[^A-Za-z0-9]+', '_', course['name'])[:60]}.pdf",
                           mime="application/pdf")


# ── PAGE: Reports ───────────────────────────────────────────────────
elif page == "📄 Reports":
    header.page_header("📄 Reports Dashboard",
                       "Every check at a glance — search, filter and export "
                       "validation reports.",
                       chip="Results and exports")
    reports = all_reports()
    visible_reports = []
    if not reports:
        st.info("No reports yet — run a check first.")
    else:
        report_ui.summary_cards(reports)
        reports = report_ui.filter_controls(reports)
        if not reports:
            st.warning("No reports match the current search/filter.")
        else:
            # Rendering hundreds of expanders and buttons on every Streamlit
            # rerun makes sidebar navigation feel slow. Keep filtering over the
            # full dataset, but render only the selected page of results.
            pc1, pc2 = st.columns([1, 1])
            page_size = pc1.selectbox(
                "Reports per page", [10, 25, 50], key="reports_page_size")
            page_count = max(1, (len(reports) + page_size - 1) // page_size)
            if not 1 <= int(st.session_state.get("reports_page_number", 1)) <= page_count:
                st.session_state["reports_page_number"] = 1
            report_page = int(pc2.number_input(
                "Page", min_value=1, max_value=page_count, value=1,
                step=1, key="reports_page_number"))
            start = (report_page - 1) * page_size
            visible_reports = reports[start:start + page_size]
            st.caption(
                f"Showing {start + 1}–{start + len(visible_reports)} of "
                f"{len(reports)} matching reports.")
    for r in visible_reports:
        icon = "🟢" if r["result"] == "Pass" else "🔴"
        moa_tag = f" · MoA: {r['moa_result']}" if r.get("moa_result") else ""
        with st.expander(f"{icon} {r['result']}{moa_tag} · {r['course_name']} "
                         f"· {r['created_at']}"):
            issues = json.loads(r["issues_json"] or "{}")
            if issues.get("summary"):
                st.markdown(f"> Entry Requirements: {issues['summary']}")
            moa_issues = json.loads(r.get("moa_issues_json") or "{}")
            if moa_issues.get("summary"):
                st.markdown(f"> Method of Assessment: {moa_issues['summary']}")
            course = get_course(r["course_id"])
            if course:
                # PDF creation can be comparatively expensive. Streamlit executes
                # the contents of closed expanders too, so generating every PDF
                # here made the Reports page slow merely by navigating to it.
                # Build only the report the user actually asks for, then keep the
                # bytes in session state for immediate download on later reruns.
                pdf_state_key = f"prepared_report_pdf_{r['id']}"
                p1, p2 = st.columns([1, 3])
                if p1.button("Prepare PDF", key=f"prep_{r['id']}"):
                    with st.spinner("Preparing PDF…"):
                        st.session_state[pdf_state_key] = build_pdf(course, r)
                if pdf_state_key in st.session_state:
                    p2.download_button(
                        "⬇️ Download PDF",
                        data=st.session_state[pdf_state_key],
                        file_name=f"ER_Report_{r['id']}.pdf",
                        mime="application/pdf",
                        key=f"dl_{r['id']}",
                    )
            if st.button("Delete report", key=f"delrep_{r['id']}"):
                delete_report(r["id"])
                st.rerun()


# ── PAGE: Content Quality Review ────────────────────────────────────
elif page == "✍️ Content Quality":
    st.markdown(QR_CSS, unsafe_allow_html=True)
    header.page_header("✍️ Content Quality Review",
                       "Paste or upload course content — grammar, articles, "
                       "sentence structure, capitalisation, proper nouns, "
                       "spelling, commas and consistency, colour-coded like a "
                       "proofreader's markup.",
                       chip="Editorial review")

    legend = "".join(
        f'<span style="background:{v["bg"]};border:1px solid {v["border"]};color:#2A2F3A;">{v["label"]}</span>'
        for v in PROOFREAD_CATEGORIES.values()
    )
    st.markdown(f'<div class="qr-legend">{legend}</div>', unsafe_allow_html=True)
    st.write("")

    src = st.radio(
        "Input", ["Paste text", "Upload file (.txt / .docx)"],
        horizontal=True, key="proofread_input_method"
    )
    text = ""
    if src == "Paste text":
        text = st.text_area(
            "Course content to review", height=220, key="proofread_pasted_text",
            placeholder="Paste the course description, overview or any page copy here…"
        )
    else:
        uploaded = st.file_uploader(
            "Upload content", type=["txt", "docx"], key="proofread_upload"
        )
        if uploaded:
            if uploaded.name.lower().endswith(".docx"):
                from docx import Document
                document = Document(io.BytesIO(uploaded.read()))
                text = "\n".join(
                    paragraph.text for paragraph in document.paragraphs
                    if paragraph.text.strip()
                )
            else:
                text = uploaded.read().decode("utf-8", errors="ignore")
            st.text_area(
                "Loaded content", text, height=180, key="proofread_loaded_content"
            )

    if st.button(
        "✍️ Review content quality", type="primary",
        disabled=not text.strip(), key="run_proofreading_review"
    ):
        if not openrouter_api_key():
            st.error(
                "No OpenRouter API key found — add OPENROUTER_API_KEY to "
                "Streamlit Secrets."
            )
        else:
            with st.spinner("Proofreading …"):
                try:
                    result = run_proofreading_review(text)
                    st.session_state["proofread_result"] = result
                    st.session_state["proofread_text"] = text
                except Exception as exc:
                    st.error(f"Review failed: {exc}")

    if "proofread_result" in st.session_state:
        result = st.session_state["proofread_result"]
        reviewed_text = st.session_state["proofread_text"]
        issues = result.get("issues", [])

        c1, c2, c3 = st.columns(3)
        qr_stat(c1, len(issues), "Issues found", "err" if issues else "ok")
        top_category = max(
            {issue.get("category") for issue in issues},
            key=lambda category: sum(
                1 for issue in issues if issue.get("category") == category
            ),
            default="—",
        )
        qr_stat(
            c2, PROOFREAD_CATEGORIES.get(top_category, {}).get("label", "—"),
            "Most common issue", "warn"
        )
        qr_stat(
            c3, f"{max(0, 100 - len(issues) * 4)}%",
            "Quality score", "info"
        )
        st.write("")

        marked_tab, fixed_tab, list_tab = st.tabs(
            ["🖍️ Marked-up text", "✅ Corrected text", "📋 Issue list"]
        )

        with marked_tab:
            if issues:
                st.markdown(
                    f'<div class="qr-paper">'
                    f'{annotate_proofreading_html(reviewed_text, issues)}</div>',
                    unsafe_allow_html=True,
                )
                st.caption(
                    "Hover a highlight to see the correction and explanation."
                )
            else:
                st.success("No issues found — this content is clean. 🎉")

        with fixed_tab:
            corrected = result.get("corrected_text", reviewed_text)
            st.text_area(
                "Corrected version (copy-ready)", corrected, height=260,
                key="proofread_corrected_text"
            )
            st.download_button(
                "⬇️ Download corrected text", corrected,
                file_name="corrected_content.txt", key="proofread_download"
            )

        with list_tab:
            if not issues:
                st.success("Nothing to list — no issues found.")
            for number, issue in enumerate(issues, start=1):
                category = issue.get("category", "grammar")
                style = PROOFREAD_CATEGORIES.get(
                    category, PROOFREAD_CATEGORIES["grammar"]
                )
                st.markdown(
                    f'<div class="issue-card" style="border-left-color:{style["border"]}">'
                    f'<div class="cat" style="color:{style["border"]}">'
                    f'#{number} · {style["label"]}</div>'
                    f'<span class="orig">{html.escape(str(issue.get("original", "")))}</span> → '
                    f'<span class="corr">{html.escape(str(issue.get("correction", "")))}</span><br>'
                    f'<small>{html.escape(str(issue.get("explanation", "")))}</small>'
                    f'</div>',
                    unsafe_allow_html=True,
                )


# ── PAGE: Course Overview Quality ───────────────────────────────────
elif page == "🧭 Course Overview Quality":
    st.markdown(QR_CSS, unsafe_allow_html=True)
    header.page_header(
        "✍️ Course Overview Quality",
        "Evaluate only the writing quality of the Course Overview section, including "
        "tone, clarity, readability, learner focus, flow and consistency.",
        chip="Separate quality review",
    )

    st.caption(
        "Choose a course from the imported Excel tracker, paste a Course Overview, "
        "or upload a document containing a Course Overview section."
    )

    legend = "".join(
        f'<span style="background:{v["bg"]};border:1px solid {v["border"]};'
        f'color:#2A2F3A;">{v["label"]}</span>'
        for v in QUALITY_CRITERIA.values()
    )
    with st.expander("View the Course Overview quality criteria"):
        st.markdown(f'<div class="qr-legend">{legend}</div>', unsafe_allow_html=True)
        st.caption(
            "Each editorial criterion receives a 0-100 score. The overall score "
            "is the simple average of all criteria; no custom percentage weights are used."
        )

    def evaluate_quality_input(body: str, context: str, source_label: str):
        if not body or not body.strip():
            st.error("No Course Overview is available to evaluate.")
            return
        if not openrouter_api_key():
            st.error(
                "No OpenRouter API key found — add OPENROUTER_API_KEY to "
                "Streamlit Secrets."
            )
            return
        with st.spinner("Evaluating Course Overview quality …"):
            try:
                result = run_quality_review(body.strip(), context)
                st.session_state["qr_result"] = result
                st.session_state["qr_text"] = body.strip()
                st.session_state["qr_context"] = context
                st.session_state["qr_source_label"] = source_label
                st.session_state["quality_corrected_output"] = result.get("corrected_text") or body.strip()
            except Exception as exc:
                st.error(f"Quality evaluation failed: {exc}")

    input_course, input_paste, input_upload = st.tabs([
        "🌐 Select course from Excel",
        "✍️ Paste content",
        "📄 Upload document",
    ])

    with input_course:
        courses = all_courses()
        if not courses:
            st.info(
                "No courses are available yet. Import the tracker Excel in "
                "📥 Upload & Specs; the course list will appear here automatically."
            )
            if st.button("Go to Upload & Specs", key="quality_go_upload"):
                st.session_state["main_navigation"] = "📥 Upload & Specs"
                st.rerun()
        else:
            levels = sorted({c["level"] for c in courses if c.get("level")})
            types = sorted({c["course_type"] for c in courses if c.get("course_type")})
            qf1, qf2, qf3 = st.columns([1, 1, 2])
            q_level = qf1.selectbox(
                "Level", ["All"] + levels, key="quality_course_level"
            )
            q_type = qf2.selectbox(
                "Type", ["All"] + types, key="quality_course_type"
            )
            q_search = qf3.text_input(
                "Search course name", key="quality_course_search",
                placeholder="Start typing a course name…",
            )
            filtered_courses = [
                c for c in courses
                if (q_level == "All" or c.get("level") == q_level)
                and (q_type == "All" or c.get("course_type") == q_type)
                and q_search.lower() in (c.get("name") or "").lower()
            ]
            if not filtered_courses:
                st.warning("No courses match the selected filters.")
            else:
                selected = st.selectbox(
                    "Course from the imported Excel tracker",
                    filtered_courses,
                    format_func=lambda c: (
                        f"{c['number']} — {c['name']}" if c.get("number") else c["name"]
                    ),
                    key="quality_selected_course",
                )
                selected_course = get_course(selected["id"])
                st.markdown(
                    f"**Course URL:** {selected_course.get('course_url') or '—'}  \n"
                    f"**Level / Type:** {selected_course.get('level') or '—'} · "
                    f"{selected_course.get('course_type') or '—'}"
                )

                if st.button(
                    "🌐 Load webpage overview",
                    type="secondary",
                    key="quality_load_course_overview",
                ):
                    with st.spinner("Reading the selected course webpage …"):
                        try:
                            overview, full_page_text = extract_page_overview(
                                selected_course["course_url"]
                            )
                            if len(overview.strip()) < 80 and openrouter_api_key():
                                try:
                                    ai_overview = ai_extract_overview(full_page_text)
                                    if len(ai_overview) > len(overview):
                                        overview = ai_overview
                                except Exception:
                                    pass
                            if not overview.strip():
                                raise RuntimeError(
                                    "No course overview or course-description section "
                                    "could be extracted from this webpage."
                                )
                            st.session_state["quality_course_text"] = overview.strip()
                            st.session_state["quality_course_editor"] = overview.strip()
                            st.session_state["quality_course_loaded_id"] = selected_course["id"]
                            st.session_state["quality_course_loaded_name"] = selected_course["name"]
                            st.rerun()
                        except Exception as exc:
                            st.error(f"Could not load the webpage overview: {exc}")

                loaded_text = st.session_state.get("quality_course_text", "")
                loaded_id = st.session_state.get("quality_course_loaded_id")
                if loaded_text and loaded_id == selected_course["id"]:
                    st.success(
                        f"Overview loaded for {st.session_state.get('quality_course_loaded_name', selected_course['name'])}."
                    )
                    if "quality_course_editor" not in st.session_state:
                        st.session_state["quality_course_editor"] = loaded_text
                    course_text = st.text_area(
                        "Extracted overview — review or edit before evaluation",
                        height=260,
                        key="quality_course_editor",
                    )
                    if st.button(
                        "🔍 Evaluate selected course overview",
                        type="primary",
                        key="quality_evaluate_course",
                    ):
                        evaluate_quality_input(
                            course_text,
                            course_quality_context(selected_course),
                            selected_course["name"],
                        )
                elif loaded_text:
                    st.info(
                        "A different course overview is currently loaded. Select that "
                        "course again or click ‘Load webpage overview’ for this course."
                    )

    with input_paste:
        paste_name = st.text_input(
            "Course name or page title (optional)",
            key="quality_paste_name",
            placeholder="e.g. NCFE CACHE Level 2 Diploma in Care",
        )
        pasted_text = st.text_area(
            "Paste the Course Overview",
            height=300,
            key="quality_pasted_text",
            placeholder="Paste the Course Overview text here…",
        )
        if st.button(
            "🔍 Evaluate pasted overview",
            type="primary",
            disabled=not pasted_text.strip(),
            key="quality_evaluate_paste",
        ):
            context = (
                f"Course/page title supplied by the user: {paste_name}"
                if paste_name.strip() else "No additional course context supplied."
            )
            overview_text = extract_course_overview_section(pasted_text)
            evaluate_quality_input(
                overview_text,
                context,
                paste_name.strip() or "Pasted Course Overview",
            )

    with input_upload:
        upload_name = st.text_input(
            "Course name or page title (optional)",
            key="quality_upload_name",
            placeholder="Used as context for the evaluation",
        )
        uploaded_doc = st.file_uploader(
            "Upload a document containing the Course Overview",
            type=["txt", "md", "docx", "pdf"],
            key="quality_document_upload",
        )
        upload_ui.info(
            ".txt · .md · .docx · .pdf",
            "PDFs must contain selectable text; scanned-image PDFs are not supported.",
        )
        uploaded_text = ""
        uploaded_overview = ""
        if uploaded_doc:
            upload_ui.file_details(uploaded_doc)
            try:
                uploaded_text = extract_uploaded_content(uploaded_doc)
                uploaded_overview = extract_course_overview_section(uploaded_text)
                upload_preview_key = (
                    "quality_uploaded_preview_"
                    + hashlib.sha1(uploaded_doc.getvalue()).hexdigest()[:10]
                )
                st.text_area(
                    "Course Overview that will be evaluated",
                    value=uploaded_overview,
                    height=260,
                    disabled=True,
                    key=upload_preview_key,
                )
            except Exception as exc:
                st.error(str(exc))
        if st.button(
            "🔍 Evaluate uploaded overview",
            type="primary",
            disabled=not uploaded_overview.strip(),
            key="quality_evaluate_upload",
        ):
            context = (
                f"Course/page title supplied by the user: {upload_name}\n"
                f"Uploaded document: {uploaded_doc.name}"
                if upload_name.strip()
                else f"Uploaded document: {uploaded_doc.name}"
            )
            evaluate_quality_input(
                uploaded_overview,
                context,
                upload_name.strip() or uploaded_doc.name,
            )

    if "qr_result" in st.session_state:
        st.divider()
        result = st.session_state["qr_result"]
        reviewed_text = st.session_state.get("qr_text", "")
        source_label = st.session_state.get("qr_source_label", "Course Overview")
        issues = result.get("issues", [])
        metrics = content_metrics(reviewed_text)

        st.subheader(f"Quality result — {source_label}")
        status_kind = {
            "Pass": "ok", "Needs Improvement": "warn", "Fail": "err"
        }.get(result.get("status"), "info")
        r1, r2, r3, r4 = st.columns(4)
        qr_stat(r1, f"{result.get('overall_score', 0)}%", "Overall quality", status_kind)
        qr_stat(r2, result.get("status", "—"), "Status", status_kind)
        priority_count = sum(
            1 for item in issues if item.get("severity") == "major"
        )
        qr_stat(r3, len(issues), "Issues identified", "err" if issues else "ok")
        qr_stat(r4, priority_count, "Priority issues", "warn" if priority_count else "ok")

        if result.get("summary"):
            st.markdown(f"> {result['summary']}")

        m1, m2, m3, m4 = st.columns(4)
        qr_stat(m1, metrics["word_count"], "Words", "info")
        qr_stat(m2, metrics["average_sentence_words"], "Avg. words / sentence", "info")
        qr_stat(m3, metrics["long_sentences_25_plus"], "Sentences 25+ words", "warn")
        qr_stat(m4, metrics["paragraph_count"], "Paragraphs", "info")

        result_overview, result_scores, result_marked, result_issues, result_rewrite = st.tabs([
            "📌 Overview",
            "📊 Category scores",
            "🖍️ Highlighted text",
            "📋 Issues",
            "✅ Improved version",
        ])

        with result_overview:
            left, right = st.columns(2)
            with left:
                st.markdown("#### Strengths")
                strengths = result.get("strengths") or []
                if strengths:
                    for strength in strengths:
                        st.markdown(f"- {strength}")
                else:
                    st.caption("No specific strengths were returned.")
            with right:
                st.markdown("#### Priority improvements")
                priority = [i for i in issues if i.get("severity") == "major"]
                priority = priority or issues[:5]
                if priority:
                    for item in priority[:6]:
                        label = QUALITY_CRITERIA.get(
                            item.get("criterion"), QUALITY_CRITERIA["language_accuracy"]
                        )["label"]
                        st.markdown(
                            f"- **{item['severity'].title()} · {label}:** "
                            f"{item.get('explanation') or item.get('correction') or 'Review this wording.'}"
                        )
                else:
                    st.success("No priority issues were identified.")

        with result_scores:
            for key, meta in QUALITY_CRITERIA.items():
                item = result.get("category_scores", {}).get(key, {})
                score = item.get("score", 0)
                explanation = html.escape(str(item.get("explanation") or ""))
                st.markdown(
                    f'<div class="quality-score-row">'
                    f'<div><div class="quality-score-label">{meta["label"]}</div>'
                    f'<div class="quality-score-explanation">{explanation}</div></div>'
                    f'<div class="quality-score-number" style="color:{meta["border"]}">{score}/100</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

        with result_marked:
            highlightable = [i for i in issues if i.get("original")]
            if highlightable:
                st.markdown(
                    f'<div class="qr-paper">{annotate_text_html(reviewed_text, highlightable)}</div>',
                    unsafe_allow_html=True,
                )
                st.caption("Hover over a highlight to see its correction and explanation.")
            elif issues:
                st.info(
                    "Issues were identified, but the AI did not return exact source "
                    "substrings that could be highlighted safely. See the Issues tab."
                )
            else:
                st.success("No text-level issues were identified.")

        with result_issues:
            if not issues:
                st.success("No issues were identified.")
            for n, issue in enumerate(issues, start=1):
                criterion = issue.get("criterion", "language_accuracy")
                style = QUALITY_CRITERIA.get(
                    criterion, QUALITY_CRITERIA["language_accuracy"]
                )
                original = html.escape(str(issue.get("original") or ""))
                correction = html.escape(str(issue.get("correction") or ""))
                comparison = ""
                if original or correction:
                    comparison = (
                        f'<span class="orig">{original or "—"}</span> → '
                        f'<span class="corr">{correction or "—"}</span><br>'
                    )
                st.markdown(
                    f'<div class="issue-card" style="border-left-color:{style["border"]}">'
                    f'<div class="cat" style="color:{style["border"]}">#{n} · '
                    f'{html.escape(issue.get("severity", "minor").title())} · {style["label"]}</div>'
                    f'{comparison}'
                    f'<small>{html.escape(str(issue.get("explanation") or ""))}</small>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

        with result_rewrite:
            corrected = result.get("corrected_text") or reviewed_text
            st.caption(
                "The improved version preserves the supplied course facts and meaning while "
                "addressing only the identified writing-quality issues."
            )
            if "quality_corrected_output" not in st.session_state:
                st.session_state["quality_corrected_output"] = corrected
            st.text_area(
                "Improved Course Overview (copy-ready)",
                height=360,
                key="quality_corrected_output",
            )
            safe_name = re.sub(r"[^A-Za-z0-9]+", "_", source_label).strip("_")[:60]
            st.download_button(
                "⬇️ Download improved text",
                st.session_state.get("quality_corrected_output", corrected),
                file_name=f"{safe_name or 'course_overview'}_improved.txt",
                mime="text/plain",
                key="quality_download_corrected",
            )


# ── PAGE: Manage Data ───────────────────────────────────────────────
else:
    header.page_header("🗂️ Manage Data",
                       "Everything you upload stays in the local database "
                       "until you remove it here.",
                       chip="Workspace administration")
    st.caption("Everything you upload is kept in the local database "
               f"(`{DB_PATH}`) until you remove it here.")

    st.subheader("Courses")
    manage_courses = all_courses()
    if manage_courses:
        mc1, mc2 = st.columns([1, 1])
        course_page_size = mc1.selectbox(
            "Courses per page", [20, 50, 100], key="manage_course_page_size")
        course_page_count = max(
            1, (len(manage_courses) + course_page_size - 1) // course_page_size)
        if not 1 <= int(st.session_state.get("manage_course_page", 1)) <= course_page_count:
            st.session_state["manage_course_page"] = 1
        course_page = int(mc2.number_input(
            "Course page", min_value=1, max_value=course_page_count,
            value=1, step=1, key="manage_course_page"))
        course_start = (course_page - 1) * course_page_size
        visible_courses = manage_courses[course_start:course_start + course_page_size]
        st.caption(
            f"Showing {course_start + 1}–{course_start + len(visible_courses)} "
            f"of {len(manage_courses)} courses.")
        for c in visible_courses:
            col1, col2 = st.columns([6, 1])
            col1.markdown(f"**{c['number']} — {c['name']}**  \n{c['course_url']}")
            if col2.button("🗑️ Remove", key=f"delc_{c['id']}"):
                delete_course(c["id"])
                st.rerun()
    else:
        st.caption("No courses stored.")

    st.subheader("Cached Specification Extractions")
    manage_specs = all_specs()
    if manage_specs:
        ms1, ms2 = st.columns([1, 1])
        spec_page_size = ms1.selectbox(
            "Specifications per page", [20, 50, 100],
            key="manage_spec_page_size")
        spec_page_count = max(
            1, (len(manage_specs) + spec_page_size - 1) // spec_page_size)
        if not 1 <= int(st.session_state.get("manage_spec_page", 1)) <= spec_page_count:
            st.session_state["manage_spec_page"] = 1
        spec_page = int(ms2.number_input(
            "Specification page", min_value=1, max_value=spec_page_count,
            value=1, step=1, key="manage_spec_page"))
        spec_start = (spec_page - 1) * spec_page_size
        visible_specs = manage_specs[spec_start:spec_start + spec_page_size]
        st.caption(
            f"Showing {spec_start + 1}–{spec_start + len(visible_specs)} "
            f"of {len(manage_specs)} specifications.")
        for s in visible_specs:
            col1, col2 = st.columns([6, 1])
            col1.markdown(f"{'🟢' if s['status'] == 'ok' else '🔴'} {s['url']}")
            if col2.button("🗑️ Remove", key=f"dels_{s['id']}"):
                delete_spec(s["id"])
                st.rerun()
    else:
        st.caption("No cached specifications stored.")

    st.divider()
    st.subheader("Danger zone")
    confirm = st.checkbox("I understand this deletes ALL courses, cached specs "
                          "and reports.")
    if st.button("🗑️ Clear all data", type="primary", disabled=not confirm):
        clear_all_data()
        st.success("All data removed.")
        st.rerun()